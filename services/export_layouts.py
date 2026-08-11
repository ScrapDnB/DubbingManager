"""Layout renderers for montage sheet HTML and DOCX exports."""

import logging
from html import escape
from typing import Any, Dict, List, Optional, Tuple

from services.assignment_service import (
    get_actor_for_character,
    get_actor_ids_for_character,
)
from services.layout_template_service import normalize_layout_template
from utils.helpers import (
    hex_to_rgba_string,
    format_seconds_to_full_tc,
    format_seconds_to_tc,
)
from utils.i18n import translate_source

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    DOCX_LAYOUT_AVAILABLE = True
except ImportError:
    DOCX_LAYOUT_AVAILABLE = False
    Document = None
    WD_ORIENT = None
    WD_CELL_VERTICAL_ALIGNMENT = None
    WD_TABLE_ALIGNMENT = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    qn = None
    Cm = None
    Pt = None
    RGBColor = None
    logger.warning("python-docx not available - DOCX layout export disabled")


class ExportLayoutMixin:
    """HTML and DOCX layout builders for montage sheets."""

    def _normalize_layout_type(self, layout_type: Optional[str]) -> str:
        """Return a supported montage sheet layout name."""
        if layout_type == "Сценарий":
            return "Сценарий 1"
        if layout_type in {
            "Таблица", "Сценарий 1", "Сценарий 2", "Сценарий 3"
        }:
            return str(layout_type)
        return "Таблица"

    @staticmethod
    def _export_font_family(cfg: Dict[str, Any]) -> str:
        """Return the selected montage-sheet font family."""
        family = str(cfg.get('font_family', 'Segoe UI') or '').strip()
        family = family.replace('\x00', '').replace('\r', ' ').replace('\n', ' ')
        return family[:100] or 'Segoe UI'

    def _table_column_width_rem(self, cfg: Dict[str, Any], key: str) -> float:
        """Return a sane table column width in rem units."""
        defaults = {
            'table_width_time': 7.0,
            'table_width_char': 10.0,
            'table_width_actor': 8.5,
        }
        try:
            value = float(cfg.get(key, defaults[key]))
        except (TypeError, ValueError):
            value = defaults[key]
        return max(4.0, min(24.0, value))

    def _table_column_width_css(self, cfg: Dict[str, Any], key: str) -> str:
        """Return responsive CSS width for a metadata table column."""
        width = self._table_column_width_rem(cfg, key)
        min_width = max(4.0, width * 0.72)
        vw_width = max(5.0, width * 1.05)
        return f"clamp({min_width:.2f}rem, {vw_width:.2f}vw, {width:.2f}rem)"

    def _table_column_width_cm(self, cfg: Dict[str, Any], key: str) -> float:
        """Return DOCX column width matching table preview proportions."""
        return self._table_column_width_rem(cfg, key) * 0.18

    def _docx_font_size_from_cfg(
        self,
        cfg: Dict[str, Any],
        key: str,
        default: float
    ) -> float:
        """Map preview pixel font settings to readable DOCX point sizes."""
        try:
            value = float(cfg.get(key, default))
        except (TypeError, ValueError):
            value = default
        return max(7.0, min(18.0, value * 0.5))

    @staticmethod
    def _export_element_bold(cfg: Dict[str, Any], key: str) -> bool:
        """Return the configured bold style for a montage sheet element."""
        setting = {
            'time': 'bold_time',
            'char': 'bold_char',
            'actor': 'bold_actor',
            'text': 'bold_text',
        }.get(key, '')
        return bool(cfg.get(setting, key == 'char')) if setting else False

    def generate_html(
        self,
        ep: str,
        processed: List[Dict[str, Any]],
        cfg: Dict[str, Any],
        highlight_ids: Optional[List[str]] = None,
        layout_type: str = "Таблица",
        is_editable: bool = True
    ) -> str:
        """Generate HTML export content."""
        layout_type = self._normalize_layout_type(layout_type)
        js = self._get_js_for_mode(is_editable)
        html = self._get_html_header(js, cfg)

        project_name = escape(str(self.project_data.get('project_name', 'Project')))
        ep_label = escape(str(ep))
        html += f"<h1>{project_name} - Серия {ep_label}</h1>"

        actors = self.project_data.get("actors", {})
        all_actor_ids = set(actors.keys())
        is_full_filter = (
            highlight_ids is not None and
            set(highlight_ids) == all_actor_ids
        )
        effective_filter = (
            None
            if (highlight_ids is None or is_full_filter)
            else set(highlight_ids)
        )
        use_color = cfg.get('use_color', True)
        soften_colors = cfg.get('soften_colors', True)
        custom_template = cfg.get('layout_template')
        if isinstance(custom_template, dict):
            custom_template = normalize_layout_template(
                custom_template, forced_kind='montage'
            )
        else:
            custom_template = None

        for idx, line in enumerate(processed):
            # Validate replica data
            if 'char' not in line:
                logger.warning(f"Skipping line without 'char' field: {line}")
                continue
            if 'text' not in line:
                logger.warning(f"Skipping line without 'text' field: {line}")
                continue

            aid, actor, is_highlighted = self._actor_display_context(
                line['char'], ep, effective_filter
            )
            h_class = "highlighted-block" if is_highlighted else ""

            bg_color, border_col = self._get_colors(
                use_color,
                is_highlighted,
                actor,
                soften_colors,
                cfg.get('color_softening_level', 1),
            )
            text_color = self._negative_text_color(
                aid,
                cfg,
                is_highlighted
            )

            text_html = self._format_text_html(line, is_editable)

            if custom_template is not None:
                html += self._build_custom_layout_row(
                    line,
                    actor,
                    text_html,
                    bg_color,
                    border_col,
                    text_color,
                    h_class,
                    cfg,
                    custom_template['root'],
                )
            elif layout_type == "Таблица":
                html += self._build_table_row(
                    line, actor, text_html, bg_color, text_color, h_class, cfg,
                    is_first=idx == 0, is_last=idx == len(processed) - 1
                )
            elif layout_type == "Сценарий 2":
                html += self._build_scenario2_row(
                    line, actor, text_html, bg_color, text_color, h_class, cfg
                )
            elif layout_type == "Сценарий 3":
                html += self._build_scenario3_row(
                    line, actor, text_html, bg_color, text_color, h_class, cfg,
                    is_first=idx == 0,
                    is_last=idx == len(processed) - 1
                )
            else:
                html += self._build_scenario_row(
                    line, actor, text_html, bg_color, border_col, text_color,
                    h_class, cfg
                )

        return html + "</body></html>"

    def _build_custom_layout_row(
        self,
        line: Dict[str, Any],
        actor: Dict[str, Any],
        text_html: str,
        bg_color: str,
        border_color: str,
        text_color: Optional[str],
        h_class: str,
        cfg: Dict[str, Any],
        root: Dict[str, Any],
    ) -> str:
        """Render a validated semantic layout tree as portable flex HTML."""
        visible = {
            'timecode': bool(cfg.get('col_tc', True)),
            'character': bool(cfg.get('col_char', True)),
            'actor': bool(cfg.get('col_actor', True)),
            'replica': bool(cfg.get('col_text', True)),
        }
        values = {
            'timecode': escape(self._format_timing_text(line, cfg)),
            'character': self._character_highlight_html(
                escape(str(line.get('char', ''))),
                bg_color,
                text_color,
                h_class,
                cfg,
            ),
            'actor': escape(str(actor.get('name', '-'))),
            'replica': text_html,
        }

        def render(node: Dict[str, Any]) -> str:
            node_type = str(node.get('type'))
            weight = max(1, min(10, int(node.get('weight', 1))))
            if node_type in {'row', 'column'}:
                children = ''.join(
                    render(child) for child in node.get('children', [])
                    if isinstance(child, dict)
                )
                if not children:
                    return ''
                direction = 'row' if node_type == 'row' else 'column'
                align = 'flex-start' if node_type == 'row' else 'stretch'
                return (
                    "<div class='custom-layout-node' style='display:flex;"
                    f"flex-direction:{direction};align-items:{align};"
                    f"gap:{int(node.get('gap', 8))}px;flex:{weight} 1 0;"
                    "min-width:0'>"
                    f"{children}</div>"
                )
            if node_type == 'separator':
                return (
                    "<div style='height:1px;background:#b8b8b8;"
                    "margin:4px 0;flex:1 1 auto'></div>"
                )
            if node_type == 'spacer':
                return (
                    f"<div style='min-height:{int(node.get('size', 12))}px'>"
                    "</div>"
                )
            field = str(node.get('field') or 'replica')
            if not visible.get(field, False):
                return ''
            style = node.get('style', {})
            font_size = max(8, min(72, int(style.get('font_size', 24))))
            font_weight = '700' if style.get('bold') else '400'
            font_style = 'italic' if style.get('italic') else 'normal'
            alignment = str(style.get('alignment', 'left'))
            if alignment not in {'left', 'center', 'right'}:
                alignment = 'left'
            return (
                f"<div class='custom-field custom-{escape(field)}' "
                f"style='flex:{weight} 1 0;min-width:0;"
                f"font-size:{font_size}px;font-weight:{font_weight};"
                f"font-style:{font_style};text-align:{alignment};"
                "white-space:pre-wrap;overflow-wrap:anywhere'>"
                f"{values.get(field, '')}</div>"
            )

        block_bg, block_border, block_text = self._container_highlight_colors(
            bg_color, border_color, text_color, h_class, cfg
        )
        return (
            f"<div class='line-container custom-layout-container {h_class}' "
            f"style='background-color:{block_bg};border-left-color:{block_border};"
            f"color:{block_text or '#000000'}'>"
            f"{render(root)}</div>"
        )

    def _actor_display_context(
        self,
        character: str,
        episode: str,
        effective_filter: Optional[set[str]],
    ) -> Tuple[Optional[str], Dict[str, Any], bool]:
        """Use actor color only when a role resolves to one selected actor."""
        actors = self.project_data.get("actors", {})
        actor_ids = get_actor_ids_for_character(
            self.project_data, character, episode
        )
        selected_ids = (
            actor_ids if effective_filter is None
            else [actor_id for actor_id in actor_ids if actor_id in effective_filter]
        )
        color_actor_id = selected_ids[0] if len(selected_ids) == 1 else None
        actor = dict(actors.get(color_actor_id, {})) if color_actor_id else {}
        actor["name"] = " / ".join(
            str(actors.get(actor_id, {}).get("name") or actor_id)
            for actor_id in actor_ids
        ) or "-"
        return color_actor_id, actor, bool(color_actor_id)

    def _get_js_for_mode(self, is_editable: bool) -> str:
        """Return js for mode."""
        if is_editable:
            return self._get_editable_js()
        return self._get_static_css()

    def _get_html_header(self, js: str, cfg: Dict[str, Any]) -> str:
        """Return the HTML document header."""
        font_family = self._export_font_family(cfg)
        css_font_family = font_family.replace('\\', '\\\\').replace("'", "\\'")
        time_width = self._table_column_width_css(cfg, 'table_width_time')
        char_width = self._table_column_width_css(cfg, 'table_width_char')
        actor_width = self._table_column_width_css(cfg, 'table_width_actor')
        time_weight = 700 if cfg.get('bold_time', False) else 400
        char_weight = 700 if cfg.get('bold_char', True) else 400
        actor_weight = 700 if cfg.get('bold_actor', False) else 400
        text_weight = 700 if cfg.get('bold_text', False) else 400
        return f"""<html><head><meta charset='utf-8'>{js}<style>
        body {{
            font-family: '{css_font_family}', sans-serif;
            padding: 36px clamp(18px, 4vw, 64px);
            background: #f6f7f8;
            color: #202124;
        }}
        h1 {{
            margin: 0 0 24px;
            font-size: 24px;
            font-weight: 650;
            line-height: 1.25;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            table-layout: fixed;
            background: white;
            border: 1px solid #d8dde3;
        }}
        td, th {{
            border: 1px solid #d8dde3;
            padding: 8px 10px;
            vertical-align: top;
            overflow-wrap: break-word;
            word-break: normal;
        }}
        .montage-table {{
            border-color: #000000;
        }}
        .montage-table td,
        .montage-table th {{
            border-color: #000000;
        }}
        th {{
            background: #eef1f4;
            color: #4f5965;
            font-size: 12px;
            font-weight: 650;
            text-align: left;
            padding: 7px 10px;
        }}
        .col-t {{
            width: {time_width};
        }}
        .col-c {{
            width: {char_width};
        }}
        .col-a {{
            width: {actor_width};
        }}
        .col-txt {{
            width: auto;
        }}
        .t {{
            font-size: {cfg.get('f_time', 12)}px;
            font-weight: {time_weight};
            line-height: 1.25;
            color: inherit;
            white-space: normal;
        }}
        .c {{
            font-weight: {char_weight};
            font-size: {cfg.get('f_char', 14)}px;
            line-height: 1.25;
        }}
        .character-highlight {{
            padding: 0.04em 0.22em 0.1em;
            border-radius: 0.12em;
            box-decoration-break: clone;
            -webkit-box-decoration-break: clone;
        }}
        .a {{
            font-style: italic;
            font-size: {cfg.get('f_actor', 14)}px;
            font-weight: {actor_weight};
            line-height: 1.25;
        }}
        .txt {{
            font-size: {cfg.get('f_text', 16)}px;
            font-weight: {text_weight};
            line-height: 1.45;
        }}
        .time-sep {{
            display: block;
            line-height: 0.9;
            opacity: 0.75;
        }}
        .line-container {{
            margin-bottom: 30px;
            padding: 20px;
            border-left: 8px solid #eee;
            background: white;
        }}
        .script2-container {{
            margin-bottom: 22px;
            padding: 16px 18px 18px;
            border: 1px solid #d8dde3;
            background: white;
        }}
        .script2-meta {{
            display: flex;
            align-items: center;
            gap: 14px;
            margin-bottom: 8px;
            color: inherit;
        }}
        .script2-time {{
            font-size: {cfg.get('f_time', 12)}px;
            font-weight: {time_weight};
            line-height: 1.25;
            white-space: pre-line;
            opacity: 0.78;
        }}
        .script2-char {{
            font-size: {cfg.get('f_char', 14)}px;
            font-weight: {char_weight};
            text-transform: uppercase;
        }}
        .script2-actor {{
            font-size: {cfg.get('f_actor', 14)}px;
            font-weight: {actor_weight};
            opacity: 0.78;
        }}
        .script2-sep {{
            opacity: 0.55;
            font-weight: 650;
        }}
        .script2-text {{
            font-size: {cfg.get('f_text', 16)}px;
            font-weight: {text_weight};
            line-height: 1.38;
        }}
        .script3-table {{
            width: 100%;
            border-collapse: collapse;
            table-layout: fixed;
            background: white;
            border: 1px solid #000000;
        }}
        .script3-table td,
        .script3-table th {{
            border-color: #000000;
        }}
        .script3-meta-col {{
            width: 26%;
        }}
        .script3-text-col {{
            width: 74%;
        }}
        .script3-meta-cell {{
            vertical-align: top;
            padding: 10px 12px;
        }}
        .script3-text-cell {{
            vertical-align: top;
            padding: 12px 14px;
        }}
        .script3-time {{
            font-size: {cfg.get('f_time', 12)}px;
            font-weight: {time_weight};
            line-height: 1.25;
            white-space: pre-line;
            opacity: 0.78;
            margin-bottom: 7px;
        }}
        .script3-char {{
            font-size: {cfg.get('f_char', 14)}px;
            font-weight: {char_weight};
            line-height: 1.2;
            text-transform: uppercase;
            margin-bottom: 4px;
        }}
        .script3-actor {{
            font-size: {cfg.get('f_actor', 14)}px;
            font-weight: {actor_weight};
            font-style: italic;
            line-height: 1.25;
            opacity: 0.78;
        }}
        .script3-text {{
            font-size: {cfg.get('f_text', 16)}px;
            font-weight: {text_weight};
            line-height: 1.38;
        }}
        </style></head><body>
        """

    def _get_editable_js(self) -> str:
        """Return editable js."""
        return """
        <script src="qrc:///qtwebchannel/qwebchannel.js"></script>
        <script>
            var backend;
            new QWebChannel(qt.webChannelTransport, function (channel) {
                backend = channel.objects.backend;
            });

            function onBlur(el) {
                if(backend) {
                    var cleanText = el.innerText;
                    cleanText = cleanText.replace(/(\\r\\n|\\n|\\r)/gm, "\\n");
                    backend.update_text(el.id, cleanText);
                }
            }
            function onKeyPress(e, el) {
                if (e.keyCode === 13) {
                    e.preventDefault();
                    el.blur();
                }
            }
        </script>
        <style>
            .edit-span {
                border-bottom: 1px dashed #ccc;
                padding: 1px 2px;
            }
            .edit-span:focus {
                background-color: #fff;
                outline: 2px solid #5B9BD5;
                border-bottom: none;
            }
            .sep {
                color: #888;
                font-weight: bold;
            }
            .highlighted-block {
                transition: outline 0.3s, box-shadow 0.3s;
            }
        </style>
        """

    def _get_static_css(self) -> str:
        """Return CSS for static HTML export."""
        return """
        <style>
            .highlighted-block {
                transition: outline 0.3s, box-shadow 0.3s;
            }
        </style>
        """

    def _get_colors(
        self,
        use_color: bool,
        is_highlighted: bool,
        actor: Dict[str, Any],
        soften_colors: bool = True,
        softening_level: int = 1,
    ) -> tuple:
        """Return colors for an export row."""
        if use_color and is_highlighted:
            actor_color = actor.get('color', '#ffffff')
            bg_color = (
                hex_to_rgba_string(
                    actor_color,
                    self._color_softening_alpha(softening_level),
                )
                if soften_colors
                else actor_color
            )
            border_col = actor_color
        else:
            bg_color = "#ffffff"
            border_col = "#eee"
        return bg_color, border_col

    @staticmethod
    def _color_softening_alpha(level: Any) -> float:
        """Map five softening levels to actor-color opacity."""
        try:
            normalized = max(-2, min(2, int(level)))
        except (TypeError, ValueError):
            normalized = 1
        return {
            -2: 0.72,
            -1: 0.55,
            0: 0.38,
            1: 0.22,
            2: 0.12,
        }[normalized]

    def _character_only_highlight(
        self,
        cfg: Dict[str, Any],
        h_class: str,
    ) -> bool:
        """Return whether this row should highlight only its character name."""
        return bool(
            cfg.get('use_color', True)
            and cfg.get('highlight_character_only', False)
            and h_class
        )

    def _character_highlight_html(
        self,
        value: str,
        bg_color: str,
        text_color: Optional[str],
        h_class: str,
        cfg: Dict[str, Any],
    ) -> str:
        """Wrap a character name in an actor-colored highlighter mark."""
        if not self._character_only_highlight(cfg, h_class):
            return value
        color_style = f";color:{text_color}" if text_color else ""
        return (
            "<span class='character-highlight' "
            f"style='background-color:{bg_color}{color_style}'>"
            f"{value}</span>"
        )

    def _container_highlight_colors(
        self,
        bg_color: str,
        border_color: str,
        text_color: Optional[str],
        h_class: str,
        cfg: Dict[str, Any],
    ) -> Tuple[str, str, Optional[str]]:
        """Remove block coloring while character-only highlighting is active."""
        if self._character_only_highlight(cfg, h_class):
            return "#ffffff", "#eee", None
        return bg_color, border_color, text_color

    def _negative_text_color(
        self,
        actor_id: Optional[str],
        cfg: Dict[str, Any],
        is_highlighted: bool
    ) -> Optional[str]:
        """Return white text when actor highlight is marked as negative."""
        if not actor_id or not is_highlighted:
            return None
        negative_ids = set(cfg.get('highlight_negative_ids_export') or [])
        if actor_id in negative_ids:
            return "#ffffff"
        return None

    def _format_text_html(
        self,
        line: Dict[str, Any],
        is_editable: bool
    ) -> str:
        """Format text html."""
        text_html = ""

        if 'parts' in line:
            for part in line['parts']:
                if part['sep']:
                    sep = escape(str(part['sep']))
                    text_html += f"<span class='sep'>{sep}</span>"
                part_text = escape(str(part.get('text', '')))
                if is_editable:
                    part_id = escape(str(part.get('id', '')), quote=True)
                    text_html += (
                        f"<span id='{part_id}' "
                        f"class='edit-span' "
                        f"contenteditable='true' "
                        f"onblur='onBlur(this)' "
                        f"onkeypress='onKeyPress(event, this)'>"
                        f"{part_text}</span>"
                    )
                else:
                    text_html += f"<span>{part_text}</span>"
        else:
            text_html = escape(str(line.get('text', '')))

        return text_html

    def _format_timing_html(self, line: Dict[str, Any], cfg: Dict[str, Any]) -> str:
        """Format timing for table HTML, allowing ranges to wrap cleanly."""
        start_tc, end_tc = self._format_timing_parts(line, cfg)
        if cfg.get('time_display', 'range') == 'start':
            return escape(start_tc)

        return (
            f"{escape(start_tc)}"
            "<span class='time-sep'>-</span>"
            f"{escape(end_tc)}"
        )

    def _format_timing_parts(
        self,
        line: Dict[str, Any],
        cfg: Dict[str, Any]
    ) -> Tuple[str, str]:
        """Return formatted start and end timing strings."""
        start = float(line.get('s', 0.0))
        end = float(line.get('e', 0.0))

        if cfg.get('round_time', False):
            start_tc = format_seconds_to_tc(start)
            end_tc = format_seconds_to_tc(end)
        else:
            start_tc = format_seconds_to_full_tc(start)
            end_tc = format_seconds_to_full_tc(end)

        if cfg.get('hide_leading_timecode_zeros', False):
            if start_tc.startswith("0:"):
                start_tc = start_tc[2:]
            if end_tc.startswith("0:"):
                end_tc = end_tc[2:]

        return start_tc, end_tc

    def _format_timing_text(self, line: Dict[str, Any], cfg: Dict[str, Any]) -> str:
        """Format timing as plain text for scenario HTML."""
        start_tc, end_tc = self._format_timing_parts(line, cfg)
        if cfg.get('time_display', 'range') == 'start':
            return start_tc
        return f"{start_tc} - {end_tc}"

    def _format_scenario2_timing_text(
        self,
        line: Dict[str, Any],
        cfg: Dict[str, Any]
    ) -> str:
        """Format timing as two lines for block scenario layout."""
        start_tc, end_tc = self._format_timing_parts(line, cfg)
        if cfg.get('time_display', 'range') == 'start':
            return start_tc
        return f"{start_tc}\n{end_tc}"

    def _build_table_row(
        self,
        line: Dict[str, Any],
        actor: Dict[str, Any],
        text_html: str,
        bg_color: str,
        text_color: Optional[str],
        h_class: str,
        cfg: Dict[str, Any],
        is_first: bool,
        is_last: bool
    ) -> str:
        """Build one table export row."""
        columns = []
        if cfg.get('col_tc', True):
            timing = self._format_timing_html(line, cfg)
            columns.append((translate_source("Время"), "t", timing))
        if cfg.get('col_char', True):
            character = self._character_highlight_html(
                escape(str(line.get('char', ''))),
                bg_color,
                text_color,
                h_class,
                cfg,
            )
            columns.append((translate_source("Персонаж"), "c", character))
        if cfg.get('col_actor', True):
            actor_name = escape(str(actor.get('name', '-')))
            columns.append((translate_source("Актер"), "a", actor_name))
        if cfg.get('col_text', True):
            columns.append((translate_source("Текст"), "txt", text_html))

        if not columns:
            columns.append((" ", "txt", ""))

        cells = "".join(
            f"<td class='{css_class}'>{value}</td>"
            for _header, css_class, value in columns
        )
        row_bg, _row_border, row_text = self._container_highlight_colors(
            bg_color, "#eee", text_color, h_class, cfg
        )
        color_style = f"; color:{row_text or '#000000'}"
        row = (
            f"<tr style='background-color:{row_bg}{color_style}' "
            f"class='{h_class}'>"
            f"{cells}</tr>"
        )

        if is_first:
            colgroup = "".join(
                f"<col class='col-{css_class}'>"
                for _header, css_class, _value in columns
            )
            headers = "".join(
                f"<th>{header}</th>"
                for header, _css_class, _value in columns
            )
            header = (
                f"<table class='montage-table'><colgroup>{colgroup}"
                "</colgroup><thead><tr>"
                f"{headers}"
                "</tr></thead><tbody>"
            )
            row = header + row

        if is_last:
            row += "</tbody></table>"

        return row

    def _build_scenario_row(
        self,
        line: Dict[str, Any],
        actor: Dict[str, Any],
        text_html: str,
        bg_color: str,
        border_col: str,
        text_color: Optional[str],
        h_class: str,
        cfg: Dict[str, Any]
    ) -> str:
        """Build scenario row."""
        meta_parts = []
        if cfg.get('col_char', True):
            char = escape(str(line.get('char', '')))
            char = self._character_highlight_html(
                char, bg_color, text_color, h_class, cfg
            )
            meta_parts.append(f"<span class='c'><b>{char}</b></span>")
        if cfg.get('col_tc', True):
            timing = escape(self._format_timing_text(line, cfg))
            meta_parts.append(f"<span class='t'>[{timing}]</span>")
        if cfg.get('col_actor', True):
            actor_name = escape(str(actor.get('name', '-')))
            meta_parts.append(f"<span class='a'><i>({actor_name})</i></span>")

        meta_html = " ".join(meta_parts)
        text_block = (
            f"<div class='txt'>{text_html}</div>"
            if cfg.get('col_text', True)
            else ""
        )
        block_bg, block_border, block_text = self._container_highlight_colors(
            bg_color, border_col, text_color, h_class, cfg
        )
        color_style = f"; color:{block_text or '#000000'}"
        return (
            f"<div class='line-container {h_class}' "
            f"style='background-color:{block_bg}; "
            f"border-left-color:{block_border}{color_style}'>"
            f"<div class='meta'>{meta_html}</div>"
            f"{text_block}</div>"
        )

    def _build_scenario2_row(
        self,
        line: Dict[str, Any],
        actor: Dict[str, Any],
        text_html: str,
        bg_color: str,
        text_color: Optional[str],
        h_class: str,
        cfg: Dict[str, Any]
    ) -> str:
        """Build the larger block-style scenario row."""
        meta_parts = []
        if cfg.get('col_tc', True):
            timing = escape(self._format_scenario2_timing_text(line, cfg))
            meta_parts.append(
                f"<span class='script2-time'>{timing}</span>"
            )
        if cfg.get('col_char', True):
            char = escape(str(line.get('char', '')))
            char = self._character_highlight_html(
                char, bg_color, text_color, h_class, cfg
            )
            meta_parts.append(
                f"<span class='script2-char'>{char}</span>"
            )
        if cfg.get('col_actor', True):
            actor_name = escape(str(actor.get('name', '-')))
            meta_parts.append(
                f"<span class='script2-actor'>{actor_name}</span>"
            )

        meta_html = "<span class='script2-sep'>|</span>".join(meta_parts)
        text_block = (
            f"<div class='script2-text'>{text_html}</div>"
            if cfg.get('col_text', True)
            else ""
        )
        block_bg, _block_border, block_text = self._container_highlight_colors(
            bg_color, "#eee", text_color, h_class, cfg
        )
        color_style = f"; color:{block_text or '#000000'}"
        return (
            f"<div class='script2-container {h_class}' "
            f"style='background-color:{block_bg}{color_style}'>"
            f"<div class='script2-meta'>{meta_html}</div>"
            f"{text_block}</div>"
        )

    def _build_scenario3_row(
        self,
        line: Dict[str, Any],
        actor: Dict[str, Any],
        text_html: str,
        bg_color: str,
        text_color: Optional[str],
        h_class: str,
        cfg: Dict[str, Any],
        is_first: bool,
        is_last: bool
    ) -> str:
        """Build one two-column scenario row."""
        meta_parts = []
        if cfg.get('col_tc', True):
            timing = escape(self._format_scenario2_timing_text(line, cfg))
            meta_parts.append(f"<div class='script3-time'>{timing}</div>")
        if cfg.get('col_char', True):
            char = escape(str(line.get('char', '')))
            char = self._character_highlight_html(
                char, bg_color, text_color, h_class, cfg
            )
            meta_parts.append(f"<div class='script3-char'>{char}</div>")
        if cfg.get('col_actor', True):
            actor_name = escape(str(actor.get('name', '-')))
            meta_parts.append(f"<div class='script3-actor'>{actor_name}</div>")
        if not meta_parts:
            meta_parts.append("&nbsp;")

        text_block = (
            f"<div class='script3-text'>{text_html}</div>"
            if cfg.get('col_text', True)
            else ""
        )
        row_bg, _row_border, row_text = self._container_highlight_colors(
            bg_color, "#eee", text_color, h_class, cfg
        )
        color_style = f"; color:{row_text or '#000000'}"
        row = (
            f"<tr class='{h_class}' style='background-color:{row_bg}{color_style}'>"
            f"<td class='script3-meta-cell'>{''.join(meta_parts)}</td>"
            f"<td class='script3-text-cell'>{text_block}</td>"
            "</tr>"
        )
        if is_first:
            row = (
                "<table class='script3-table'>"
                "<colgroup>"
                "<col class='script3-meta-col'>"
                "<col class='script3-text-col'>"
                "</colgroup>"
                "<thead><tr>"
                f"<th>{translate_source('Кто / когда')}</th>"
                f"<th>{translate_source('Реплика')}</th>"
                "</tr></thead><tbody>"
            ) + row
        if is_last:
            row += "</tbody></table>"
        return row

    def _set_docx_cell_shading(self, cell: Any, color_hex: str) -> None:
        """Set a table cell background color."""
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn('w:shd'))
        if shading is None:
            shading = OxmlElement('w:shd')
            tc_pr.append(shading)
        shading.set(qn('w:fill'), color_hex.replace('#', '').upper())

    def _set_docx_run_shading(self, run: Any, color_hex: str) -> None:
        """Apply a highlighter-style background to one DOCX text run."""
        run_properties = run._element.get_or_add_rPr()
        shading = run_properties.find(qn('w:shd'))
        if shading is None:
            shading = OxmlElement('w:shd')
            run_properties.append(shading)
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), color_hex.replace('#', '').upper())

    def _set_docx_cell_margins(
        self,
        cell: Any,
        top: int = 110,
        start: int = 170,
        bottom: int = 130,
        end: int = 170
    ) -> None:
        """Set DOCX table cell internal margins in twips."""
        tc_pr = cell._tc.get_or_add_tcPr()
        margins = tc_pr.find(qn('w:tcMar'))
        if margins is None:
            margins = OxmlElement('w:tcMar')
            tc_pr.append(margins)
        for side, value in {
            'top': top,
            'start': start,
            'bottom': bottom,
            'end': end,
        }.items():
            node = margins.find(qn(f'w:{side}'))
            if node is None:
                node = OxmlElement(f'w:{side}')
                margins.append(node)
            node.set(qn('w:w'), str(value))
            node.set(qn('w:type'), 'dxa')

    def _set_docx_cell_border(
        self,
        cell: Any,
        side: str,
        color: str,
        size: int = 4
    ) -> None:
        """Set one DOCX cell border."""
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn('w:tcBorders'))
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tc_pr.append(borders)
        border = borders.find(qn(f'w:{side}'))
        if border is None:
            border = OxmlElement(f'w:{side}')
            borders.append(border)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color.replace('#', '').upper())

    def _set_docx_block_borders(
        self,
        cell: Any,
        border_color: str = "DADCE0",
        left_color: Optional[str] = None
    ) -> None:
        """Set quiet borders for a one-cell DOCX script block."""
        for side in ['top', 'right', 'bottom']:
            self._set_docx_cell_border(cell, side, border_color, 4)
        self._set_docx_cell_border(
            cell,
            'left',
            left_color or border_color,
            18 if left_color else 4
        )

    def _set_docx_cell_width(self, cell: Any, width_cm: float) -> None:
        """Set a table cell width in twentieths of a point."""
        tc_pr = cell._tc.get_or_add_tcPr()
        width = tc_pr.find(qn('w:tcW'))
        if width is None:
            width = OxmlElement('w:tcW')
            tc_pr.append(width)
        width.set(qn('w:w'), str(int(width_cm * 567)))
        width.set(qn('w:type'), 'dxa')

    def _cm_to_twips(self, width_cm: float) -> int:
        """Convert centimeters to Word twips."""
        return int(width_cm * 567)

    def _set_docx_table_grid(
        self,
        table: Any,
        column_widths: List[float]
    ) -> None:
        """Set fixed DOCX table layout and grid for better Pages support."""
        table_width = sum(column_widths)
        table_pr = table._tbl.tblPr

        layout = table_pr.find(qn('w:tblLayout'))
        if layout is None:
            layout = OxmlElement('w:tblLayout')
            table_pr.append(layout)
        layout.set(qn('w:type'), 'fixed')

        width = table_pr.find(qn('w:tblW'))
        if width is None:
            width = OxmlElement('w:tblW')
            table_pr.append(width)
        width.set(qn('w:w'), str(self._cm_to_twips(table_width)))
        width.set(qn('w:type'), 'dxa')

        existing_grid = table._tbl.tblGrid
        if existing_grid is not None:
            table._tbl.remove(existing_grid)
        grid = OxmlElement('w:tblGrid')
        for column_width in column_widths:
            grid_col = OxmlElement('w:gridCol')
            grid_col.set(qn('w:w'), str(self._cm_to_twips(column_width)))
            grid.append(grid_col)
        table._tbl.insert(0, grid)

    def _docx_soft_fill_color(self, color_hex: str, alpha: float = 0.22) -> str:
        """Return a light color similar to transparent HTML row highlight."""
        value = str(color_hex or "").strip().lstrip("#")
        if len(value) != 6:
            return "FFFFFF"
        try:
            red = int(value[0:2], 16)
            green = int(value[2:4], 16)
            blue = int(value[4:6], 16)
        except ValueError:
            return "FFFFFF"
        red = round(255 - (255 - red) * alpha)
        green = round(255 - (255 - green) * alpha)
        blue = round(255 - (255 - blue) * alpha)
        return f"{red:02X}{green:02X}{blue:02X}"

    def _set_docx_cell_text(
        self,
        cell: Any,
        text: Any,
        font_size: float,
        bold: bool = False,
        align: Any = None,
        fill_color: Optional[str] = None,
        text_color: Optional[str] = None,
        run_fill_color: Optional[str] = None,
    ) -> Any:
        """Fill a DOCX table cell with styled text."""
        if align is None:
            align = WD_ALIGN_PARAGRAPH.LEFT
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        run = paragraph.add_run(str(text or ""))
        run.bold = bold
        font_family = getattr(self, '_active_export_font_family', 'Segoe UI')
        run.font.name = font_family
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)
        run.font.size = Pt(font_size)
        if text_color:
            clean_color = text_color.strip().lstrip("#")
            if len(clean_color) == 6:
                run.font.color.rgb = RGBColor.from_string(clean_color.upper())
        if run_fill_color:
            self._set_docx_run_shading(run, run_fill_color)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        if fill_color:
            self._set_docx_cell_shading(cell, fill_color)
        return run

    def _add_docx_run(
        self,
        paragraph: Any,
        text: Any,
        font_size: float,
        bold: bool = False,
        italic: bool = False,
        color: Optional[str] = None,
        highlight_fill: Optional[str] = None,
    ) -> Any:
        """Add a consistently styled DOCX run."""
        run = paragraph.add_run(str(text or ""))
        run.bold = bold
        run.italic = italic
        font_family = getattr(self, '_active_export_font_family', 'Segoe UI')
        run.font.name = font_family
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)
        run.font.size = Pt(font_size)
        if color:
            clean_color = color.strip().lstrip("#")
            if len(clean_color) == 6:
                run.font.color.rgb = RGBColor.from_string(clean_color.upper())
        if highlight_fill:
            self._set_docx_run_shading(run, highlight_fill)
        return run

    def _create_docx_episode_table(
        self,
        document: Any,
        ep_num: str,
        processed: List[Dict[str, Any]],
        cfg: Dict[str, Any]
    ) -> None:
        """Create a DOCX table for one episode."""
        actors = self.project_data.get('actors', {})
        use_color = cfg.get('use_color', True)
        soften_colors = cfg.get('soften_colors', True)
        effective_filter = self._get_effective_highlight_filter(cfg)

        columns: List[Tuple[str, str, float, float]] = []
        if cfg.get('col_tc', True):
            columns.append((
                translate_source('Время'),
                'time',
                self._table_column_width_cm(cfg, 'table_width_time'),
                self._docx_font_size_from_cfg(cfg, 'f_time', 21),
            ))
        if cfg.get('col_char', True):
            columns.append((
                translate_source('Персонаж'),
                'char',
                self._table_column_width_cm(cfg, 'table_width_char'),
                self._docx_font_size_from_cfg(cfg, 'f_char', 20),
            ))
        if cfg.get('col_actor', True):
            columns.append((
                translate_source('Актёр'),
                'actor',
                self._table_column_width_cm(cfg, 'table_width_actor'),
                self._docx_font_size_from_cfg(cfg, 'f_actor', 14),
            ))
        if cfg.get('col_text', True):
            metadata_width = sum(column[2] for column in columns)
            columns.append((
                translate_source('Реплика'),
                'text',
                max(7.0, 18.0 - metadata_width),
                self._docx_font_size_from_cfg(cfg, 'f_text', 30),
            ))
        if not columns:
            columns.append((translate_source('Реплика'), 'text', 18.0, 11.0))

        if len(document.paragraphs) > 1 or document.tables:
            document.add_page_break()
        heading = document.add_heading(
            f"{self.project_data.get('project_name', 'Project')} - "
            f"{translate_source('Серия')} {ep_num}",
            level=1
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        table = document.add_table(rows=1, cols=len(columns))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.style = 'Table Grid'
        self._set_docx_table_grid(
            table,
            [width for _title, _key, width, _size in columns]
        )

        header_cells = table.rows[0].cells
        for cell, (title, _key, width, _size) in zip(header_cells, columns):
            self._set_docx_cell_width(cell, width)
            self._set_docx_cell_text(
                cell,
                title,
                font_size=9.0,
                bold=True,
                fill_color='EAEAEA'
            )

        for line in processed:
            row_cells = table.add_row().cells
            char_name = line.get('char', '')
            actor_id = get_actor_for_character(
                self.project_data, char_name, ep_num
            )
            actor = actors.get(actor_id, {}) if actor_id else {}
            actor_name = actor.get('name', '-') if actor else '-'

            is_highlighted = (
                effective_filter is None or
                actor_id in effective_filter
            )
            fill_color = None
            if use_color and actor_id and is_highlighted:
                actor_color = actor.get('color', '#FFFFFF')
                fill_color = (
                    self._docx_soft_fill_color(
                        actor_color,
                        self._color_softening_alpha(
                            cfg.get('color_softening_level', 1)
                        ),
                    )
                    if soften_colors
                    else actor_color.replace('#', '')
                )
            text_color = self._negative_text_color(
                actor_id,
                cfg,
                is_highlighted
            )
            character_only = bool(
                cfg.get('highlight_character_only', False) and fill_color
            )

            values = {
                'time': self._format_table_timing_text(line, cfg),
                'char': char_name,
                'actor': actor_name,
                'text': line.get('text', ''),
            }
            for cell, (_title, key, width, font_size) in zip(row_cells, columns):
                self._set_docx_cell_width(cell, width)
                self._set_docx_cell_text(
                    cell,
                    values.get(key, ''),
                    font_size=font_size,
                    bold=self._export_element_bold(cfg, key),
                    fill_color=None if character_only else fill_color,
                    text_color=(
                        text_color
                        if not character_only or key == 'char'
                        else None
                    ),
                    run_fill_color=(
                        fill_color if character_only and key == 'char' else None
                    ),
                )

    def _docx_episode_heading(self, document: Any, ep_num: str) -> None:
        """Add a DOCX episode heading, separating episodes when needed."""
        if len(document.paragraphs) > 1 or document.tables:
            document.add_page_break()
        heading = document.add_heading(
            f"{self.project_data.get('project_name', 'Project')} - "
            f"{translate_source('Серия')} {ep_num}",
            level=1
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _docx_line_actor_context(
        self,
        ep_num: str,
        line: Dict[str, Any],
        cfg: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any], bool, Optional[str], Optional[str]]:
        """Return actor context and colors for one DOCX line."""
        actors = self.project_data.get('actors', {})
        effective_filter = self._get_effective_highlight_filter(cfg)
        char_name = line.get('char', '')
        actor_id = get_actor_for_character(self.project_data, char_name, ep_num)
        actor = actors.get(actor_id, {}) if actor_id else {}
        is_highlighted = (
            effective_filter is None or
            actor_id in effective_filter
        )
        fill_color = None
        if cfg.get('use_color', True) and actor_id and is_highlighted:
            actor_color = actor.get('color', '#FFFFFF')
            fill_color = (
                self._docx_soft_fill_color(
                    actor_color,
                    self._color_softening_alpha(
                        cfg.get('color_softening_level', 1)
                    ),
                )
                if cfg.get('soften_colors', True)
                else actor_color.replace('#', '')
            )
        text_color = self._negative_text_color(actor_id, cfg, is_highlighted)
        return actor_id or "", actor, is_highlighted, fill_color, text_color

    def _add_docx_script_block(
        self,
        document: Any,
        fill_color: Optional[str],
        left_border: Optional[str] = None
    ) -> Any:
        """Create a one-cell table used as a scenario block."""
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        self._set_docx_table_grid(table, [18.0])
        cell = table.cell(0, 0)
        self._set_docx_cell_width(cell, 18.0)
        self._set_docx_cell_margins(cell)
        self._set_docx_block_borders(
            cell,
            border_color="DADCE0",
            left_color=left_border
        )
        if fill_color:
            self._set_docx_cell_shading(cell, fill_color)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        return cell

    def _create_docx_episode_scenario1(
        self,
        document: Any,
        ep_num: str,
        processed: List[Dict[str, Any]],
        cfg: Dict[str, Any]
    ) -> None:
        """Create a DOCX version of the classic HTML scenario layout."""
        self._docx_episode_heading(document, ep_num)
        for line in processed:
            _actor_id, actor, _is_highlighted, fill_color, text_color = (
                self._docx_line_actor_context(ep_num, line, cfg)
            )
            character_only = bool(
                cfg.get('highlight_character_only', False) and fill_color
            )
            left_border = None
            if fill_color and not character_only:
                left_border = actor.get('color', '#DADCE0').replace('#', '')
            run_color = "#000000" if character_only else text_color or "#000000"
            cell = self._add_docx_script_block(
                document,
                None if character_only else fill_color,
                left_border=left_border
            )

            meta = cell.paragraphs[0]
            meta.paragraph_format.space_after = Pt(4)
            meta.paragraph_format.line_spacing = 1.1
            meta_parts_added = 0
            if cfg.get('col_char', True):
                self._add_docx_run(
                    meta,
                    line.get('char', ''),
                    self._docx_font_size_from_cfg(cfg, 'f_char', 20),
                    bold=self._export_element_bold(cfg, 'char'),
                    color=(text_color or "#000000") if character_only else run_color,
                    highlight_fill=fill_color if character_only else None,
                )
                meta_parts_added += 1
            if cfg.get('col_tc', True):
                if meta_parts_added:
                    self._add_docx_run(meta, " ", 8, color=run_color)
                self._add_docx_run(
                    meta,
                    f"[{self._format_timing_text(line, cfg)}]",
                    self._docx_font_size_from_cfg(cfg, 'f_time', 21),
                    bold=self._export_element_bold(cfg, 'time'),
                    color=run_color
                )
                meta_parts_added += 1
            if cfg.get('col_actor', True):
                if meta_parts_added:
                    self._add_docx_run(meta, " ", 8, color=run_color)
                actor_name = actor.get('name', '-') if actor else '-'
                self._add_docx_run(
                    meta,
                    f"({actor_name})",
                    self._docx_font_size_from_cfg(cfg, 'f_actor', 14),
                    bold=self._export_element_bold(cfg, 'actor'),
                    italic=True,
                    color=run_color
                )

            if cfg.get('col_text', True):
                body = cell.add_paragraph()
                body.paragraph_format.space_before = Pt(0)
                body.paragraph_format.space_after = Pt(0)
                body.paragraph_format.line_spacing = 1.2
                self._add_docx_run(
                    body,
                    line.get('text', ''),
                    self._docx_font_size_from_cfg(cfg, 'f_text', 30),
                    bold=self._export_element_bold(cfg, 'text'),
                    color=run_color
                )
            spacer = document.add_paragraph()
            spacer.paragraph_format.space_after = Pt(4)

    def _create_docx_episode_scenario2(
        self,
        document: Any,
        ep_num: str,
        processed: List[Dict[str, Any]],
        cfg: Dict[str, Any]
    ) -> None:
        """Create the larger block-style DOCX scenario layout."""
        self._docx_episode_heading(document, ep_num)
        for line in processed:
            _actor_id, actor, _is_highlighted, fill_color, text_color = (
                self._docx_line_actor_context(ep_num, line, cfg)
            )
            character_only = bool(
                cfg.get('highlight_character_only', False) and fill_color
            )
            run_color = "#000000" if character_only else text_color or "#000000"
            muted_color = "#505050" if character_only else text_color or "#505050"
            cell = self._add_docx_script_block(
                document,
                None if character_only else fill_color,
            )

            meta = cell.paragraphs[0]
            meta.paragraph_format.space_after = Pt(4)
            meta.paragraph_format.line_spacing = 1.05
            meta_parts_added = 0
            if cfg.get('col_tc', True):
                self._add_docx_run(
                    meta,
                    self._format_scenario2_timing_text(line, cfg),
                    self._docx_font_size_from_cfg(cfg, 'f_time', 21),
                    bold=self._export_element_bold(cfg, 'time'),
                    color=muted_color
                )
                meta_parts_added += 1
            if cfg.get('col_char', True):
                if meta_parts_added:
                    self._add_docx_run(meta, "   |   ", 9, color=muted_color)
                self._add_docx_run(
                    meta,
                    str(line.get('char', '')).upper(),
                    self._docx_font_size_from_cfg(cfg, 'f_char', 20),
                    bold=self._export_element_bold(cfg, 'char'),
                    color=(text_color or "#000000") if character_only else run_color,
                    highlight_fill=fill_color if character_only else None,
                )
                meta_parts_added += 1
            if cfg.get('col_actor', True):
                if meta_parts_added:
                    self._add_docx_run(meta, "   |   ", 9, color=muted_color)
                actor_name = actor.get('name', '-') if actor else '-'
                self._add_docx_run(
                    meta,
                    actor_name,
                    self._docx_font_size_from_cfg(cfg, 'f_actor', 14),
                    bold=self._export_element_bold(cfg, 'actor'),
                    color=muted_color
                )

            if cfg.get('col_text', True):
                body = cell.add_paragraph()
                body.paragraph_format.space_before = Pt(1)
                body.paragraph_format.space_after = Pt(0)
                body.paragraph_format.line_spacing = 1.18
                self._add_docx_run(
                    body,
                    line.get('text', ''),
                    self._docx_font_size_from_cfg(cfg, 'f_text', 30),
                    bold=self._export_element_bold(cfg, 'text'),
                    color=run_color
                )
            spacer = document.add_paragraph()
            spacer.paragraph_format.space_after = Pt(3)

    def _create_docx_episode_scenario3(
        self,
        document: Any,
        ep_num: str,
        processed: List[Dict[str, Any]],
        cfg: Dict[str, Any]
    ) -> None:
        """Create a two-column DOCX scenario layout."""
        self._docx_episode_heading(document, ep_num)
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.style = 'Table Grid'
        column_widths = [4.1, 13.9]
        self._set_docx_table_grid(table, column_widths)

        for cell, title, width in zip(
            table.rows[0].cells,
            [translate_source('Кто / когда'), translate_source('Реплика')],
            column_widths
        ):
            self._set_docx_cell_width(cell, width)
            self._set_docx_cell_margins(cell, top=85, bottom=85)
            for side in ('top', 'right', 'bottom', 'left'):
                self._set_docx_cell_border(cell, side, '000000', 4)
            self._set_docx_cell_text(
                cell,
                title,
                font_size=9.0,
                bold=True,
                fill_color='EEF1F4',
                text_color='#4F5965'
            )

        for line in processed:
            _actor_id, actor, _is_highlighted, fill_color, text_color = (
                self._docx_line_actor_context(ep_num, line, cfg)
            )
            character_only = bool(
                cfg.get('highlight_character_only', False) and fill_color
            )
            run_color = "#000000" if character_only else text_color or "#000000"
            muted_color = "#505050" if character_only else text_color or "#505050"
            row_cells = table.add_row().cells
            for cell, width in zip(row_cells, column_widths):
                self._set_docx_cell_width(cell, width)
                self._set_docx_cell_margins(cell)
                for side in ('top', 'right', 'bottom', 'left'):
                    self._set_docx_cell_border(cell, side, '000000', 4)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                if fill_color and not character_only:
                    self._set_docx_cell_shading(cell, fill_color)

            meta_cell, text_cell = row_cells
            meta_cell.text = ""
            current_paragraph = meta_cell.paragraphs[0]
            current_paragraph.paragraph_format.space_after = Pt(3)
            current_paragraph.paragraph_format.line_spacing = 1.05
            wrote_meta = False
            if cfg.get('col_tc', True):
                self._add_docx_run(
                    current_paragraph,
                    self._format_scenario2_timing_text(line, cfg),
                    self._docx_font_size_from_cfg(cfg, 'f_time', 21),
                    bold=self._export_element_bold(cfg, 'time'),
                    color=muted_color
                )
                wrote_meta = True
            if cfg.get('col_char', True):
                current_paragraph = (
                    meta_cell.add_paragraph()
                    if wrote_meta
                    else current_paragraph
                )
                current_paragraph.paragraph_format.space_after = Pt(2)
                self._add_docx_run(
                    current_paragraph,
                    str(line.get('char', '')).upper(),
                    self._docx_font_size_from_cfg(cfg, 'f_char', 20),
                    bold=self._export_element_bold(cfg, 'char'),
                    color=(text_color or "#000000") if character_only else run_color,
                    highlight_fill=fill_color if character_only else None,
                )
                wrote_meta = True
            if cfg.get('col_actor', True):
                current_paragraph = (
                    meta_cell.add_paragraph()
                    if wrote_meta
                    else current_paragraph
                )
                current_paragraph.paragraph_format.space_after = Pt(0)
                self._add_docx_run(
                    current_paragraph,
                    actor.get('name', '-') if actor else '-',
                    self._docx_font_size_from_cfg(cfg, 'f_actor', 14),
                    bold=self._export_element_bold(cfg, 'actor'),
                    italic=True,
                    color=muted_color
                )

            text_cell.text = ""
            if cfg.get('col_text', True):
                body = text_cell.paragraphs[0]
                body.paragraph_format.space_after = Pt(0)
                body.paragraph_format.line_spacing = 1.18
                self._add_docx_run(
                    body,
                    line.get('text', ''),
                    self._docx_font_size_from_cfg(cfg, 'f_text', 30),
                    bold=self._export_element_bold(cfg, 'text'),
                    color=run_color
                )

    def create_docx_document(
        self,
        episodes_data: Dict[str, List[Dict[str, Any]]],
        cfg: Optional[Dict[str, Any]] = None
    ) -> Any:
        """Create a DOCX document with one table per episode."""
        if not DOCX_LAYOUT_AVAILABLE:
            raise ImportError("python-docx not available")

        if cfg is None:
            cfg = self.project_data.get("export_config", {})

        self._active_export_font_family = self._export_font_family(cfg)

        document = Document()
        for style_name in ('Normal', 'Title', 'Heading 1'):
            style = document.styles[style_name]
            style.font.name = self._active_export_font_family
            style._element.rPr.rFonts.set(
                qn('w:eastAsia'), self._active_export_font_family
            )
        section = document.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)

        layout_type = self._normalize_layout_type(cfg.get('layout_type'))
        custom_template = cfg.get('layout_template')
        if isinstance(custom_template, dict):
            custom_template = normalize_layout_template(
                custom_template, forced_kind='montage'
            )
        else:
            custom_template = None
        for ep_num in sorted(episodes_data.keys(), key=self._episode_sort_key):
            if custom_template is not None:
                self._create_docx_episode_custom(
                    document,
                    ep_num,
                    episodes_data[ep_num],
                    cfg,
                    custom_template['root'],
                )
            elif layout_type == "Сценарий 1":
                self._create_docx_episode_scenario1(
                    document, ep_num, episodes_data[ep_num], cfg
                )
            elif layout_type == "Сценарий 2":
                self._create_docx_episode_scenario2(
                    document, ep_num, episodes_data[ep_num], cfg
                )
            elif layout_type == "Сценарий 3":
                self._create_docx_episode_scenario3(
                    document, ep_num, episodes_data[ep_num], cfg
                )
            else:
                self._create_docx_episode_table(
                    document,
                    ep_num,
                    episodes_data[ep_num],
                    cfg
                )

        return document

    def _create_docx_episode_custom(
        self,
        document: Any,
        ep_num: str,
        lines: List[Dict[str, Any]],
        cfg: Dict[str, Any],
        root: Dict[str, Any],
    ) -> None:
        """Render the semantic tree with nested Word tables and paragraphs."""
        document.add_heading(
            f"{self.project_data.get('project_name', 'Project')} - Серия {ep_num}",
            level=1,
        )
        for line in lines:
            _actor_id, actor, _highlighted = self._actor_display_context(
                str(line.get('char', '')), str(ep_num), None
            )
            outer = document.add_table(rows=1, cols=1)
            outer.alignment = WD_TABLE_ALIGNMENT.CENTER
            outer.autofit = True
            cell = outer.cell(0, 0)
            cell.text = ""
            self._render_docx_custom_node(cell, root, line, actor, cfg)
            document.add_paragraph().paragraph_format.space_after = Pt(2)

    def _render_docx_custom_node(
        self,
        cell: Any,
        node: Dict[str, Any],
        line: Dict[str, Any],
        actor: Dict[str, Any],
        cfg: Dict[str, Any],
    ) -> None:
        node_type = str(node.get('type'))
        if node_type == 'row':
            children = [
                child for child in node.get('children', [])
                if isinstance(child, dict)
            ]
            if not children:
                return
            table = cell.add_table(rows=1, cols=len(children))
            table.autofit = True
            weights = [max(1, int(child.get('weight', 1))) for child in children]
            total_weight = max(1, sum(weights))
            for index, child in enumerate(children):
                child_cell = table.cell(0, index)
                child_cell.text = ""
                child_cell.width = Cm(17.0 * weights[index] / total_weight)
                self._render_docx_custom_node(
                    child_cell, child, line, actor, cfg
                )
            return
        if node_type == 'column':
            for child in node.get('children', []):
                if isinstance(child, dict):
                    self._render_docx_custom_node(cell, child, line, actor, cfg)
            return
        if node_type == 'separator':
            paragraph = cell.add_paragraph("────────────────")
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            return
        if node_type == 'spacer':
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(
                max(1, min(50, int(node.get('size', 12)))) * 0.25
            )
            return

        field = str(node.get('field') or 'replica')
        if field == 'timecode':
            value = self._format_timing_text(line, cfg)
        elif field == 'character':
            value = str(line.get('char', ''))
        elif field == 'actor':
            value = str(actor.get('name', '-'))
        else:
            value = str(line.get('text', ''))
        style = node.get('style', {})
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        alignment = str(style.get('alignment', 'left'))
        paragraph.alignment = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self._add_docx_run(
            paragraph,
            value,
            max(7.0, min(36.0, float(style.get('font_size', 24)) * 0.5)),
            bold=bool(style.get('bold')),
            italic=bool(style.get('italic')),
        )
