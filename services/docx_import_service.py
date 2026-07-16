"""Service for importing DOCX files with flexible column mapping."""

import os
import re
import logging
from copy import deepcopy
from typing import Dict, List, Any, Optional, Tuple
from collections import defaultdict

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("python-docx not installed. DOCX import disabled.")

from utils.helpers import srt_time_to_seconds
from config.constants import DEFAULT_DOCX_IMPORT_CONFIG

logger = logging.getLogger(__name__)


# Column types
COLUMN_TYPES = {
    'character': 'Имя персонажа',
    'time_start': 'Тайминг (начало)',
    'time_end': 'Тайминг (конец)',
    'time_split': 'Тайминг (вместе)',
    'text': 'Текст фразы'
}

# Default values for column mapping
DEFAULT_COLUMN_MAPPING = {
    'character': 0,
    'time_start': None,
    'time_end': None,
    'time_split': 1,
    'text': 2
}

# Default separators for combined timing columns
DEFAULT_TIME_SEPARATORS = list(DEFAULT_DOCX_IMPORT_CONFIG['time_separators'])


class DocxImportService:
    """Docx Import Service implementation."""

    def __init__(
        self,
        merge_gap: int = 5,
        fps: float = 25.0,
        time_separators: Optional[List[str]] = None,
        detection_config: Optional[Dict[str, Any]] = None,
    ):
        self.merge_gap = merge_gap
        self.fps = fps
        self.detection_config = deepcopy(DEFAULT_DOCX_IMPORT_CONFIG)
        if isinstance(detection_config, dict):
            self.detection_config.update({
                key: deepcopy(value)
                for key, value in detection_config.items()
                if key != 'aliases'
            })
            if isinstance(detection_config.get('aliases'), dict):
                self.detection_config['aliases'].update(
                    deepcopy(detection_config['aliases'])
                )
        self.time_separators = (
            time_separators
            or self.detection_config.get('time_separators')
            or DEFAULT_TIME_SEPARATORS
        )
        self.last_detection: Dict[str, Any] = {}

    def set_merge_gap(self, gap: int) -> None:
        """Set the replica merge gap."""
        self.merge_gap = gap

    def set_fps(self, fps: float) -> None:
        """Set the frame rate."""
        self.fps = fps

    def set_time_separators(self, separators: List[str]) -> None:
        """Set time separators."""
        self.time_separators = separators

    def extract_tables_from_docx(self, path: str) -> List[List[List[str]]]:
        """Extract tables from docx."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed")
            return []

        # Reset the header flag
        self._has_header = False

        try:
            doc = Document(path)
            all_tables = []

            # Extract all tables
            if doc.tables:
                for table in doc.tables:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if any(cells):  # Skip empty rows
                            rows.append(cells)
                    if rows:  # Add only non-empty tables
                        all_tables.append(rows)
            else:
                # If there are no tables, try parsing text as a table
                # (separators are tabs or repeated spaces)
                rows = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        # Try splitting by tabs
                        if '\t' in text:
                            cells = [c.strip() for c in text.split('\t')]
                        else:
                            cells = [text]
                        rows.append(cells)
                if rows:
                    all_tables.append(rows)

            return all_tables

        except Exception as e:
            logger.error(f"Error extracting tables from DOCX: {e}")
            return []

    def extract_first_table(self, path: str) -> List[List[str]]:
        """Extract first table."""
        all_tables = self.extract_tables_from_docx(path)
        if all_tables:
            self._has_header = False  # Reset for the first table
            return all_tables[0]
        return []

    def parse_document(
        self,
        path: str,
        column_mapping: Optional[Dict[str, Optional[int]]] = None,
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Parse every table in a DOCX using saved mapping when it still fits."""
        return self.parse_tables(
            self.extract_tables_from_docx(path),
            column_mapping,
        )

    def parse_tables(
        self,
        tables: List[List[List[str]]],
        column_mapping: Optional[Dict[str, Optional[int]]] = None,
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Parse DOCX table rows and combine them into one episode payload."""
        all_lines: List[Dict[str, Any]] = []
        for rows in tables:
            if not rows:
                continue
            detected_mapping = self.detect_columns(rows)
            column_count = max((len(row) for row in rows), default=0)
            mapping = (
                column_mapping
                if self.mapping_usable(column_mapping, column_count)
                else detected_mapping
            )
            _stats, lines = self.parse_with_mapping(rows, mapping)
            all_lines.extend(lines)
        return self._summarize_lines(all_lines), all_lines

    @staticmethod
    def mapping_usable(mapping: Any, column_count: int) -> bool:
        """Return whether a saved mapping can address the current table."""
        if not isinstance(mapping, dict) or mapping.get('text') is None:
            return False
        for field in COLUMN_TYPES:
            value = mapping.get(field)
            if value is None:
                continue
            try:
                index = int(value)
            except (TypeError, ValueError):
                return False
            if not 0 <= index < column_count:
                return False
        return True

    @staticmethod
    def _summarize_lines(
        lines: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        characters: Dict[str, Dict[str, int]] = defaultdict(
            lambda: {'lines': 0, 'words': 0}
        )
        for line in lines:
            character = str(line.get('char') or '')
            if not character:
                continue
            characters[character]['lines'] += 1
            characters[character]['words'] += len(
                str(line.get('text') or '').split()
            )
        return [
            {
                'name': character,
                'lines': values['lines'],
                'rings': values['lines'],
                'words': values['words'],
            }
            for character, values in characters.items()
        ]

    def detect_columns(self, rows: List[List[str]]) -> Dict[str, int]:
        """Detect columns."""
        if not rows:
            return DEFAULT_COLUMN_MAPPING.copy()

        config = self.detection_config
        mode = config.get('header_mode', 'auto')
        search_count = min(
            len(rows), max(1, int(config.get('header_search_rows', 5)))
        )
        candidates = [0] if mode == 'first' else range(search_count)
        best_index, best_mapping, best_score = None, {}, -1
        if mode != 'none':
            for row_index in candidates:
                mapping = self._detect_header_mapping(rows[row_index])
                score = len(mapping)
                if score > best_score:
                    best_index, best_mapping, best_score = (
                        row_index, mapping, score
                    )

        minimum = int(config.get('minimum_header_matches', 2))
        has_header = best_score >= minimum
        mapping = best_mapping if has_header else {}
        fallback = config.get('fallback_mapping') or DEFAULT_COLUMN_MAPPING
        for col_type, default_idx in fallback.items():
            if col_type not in mapping:
                if col_type == 'text':
                    max_cols = max(len(row) for row in rows)
                    mapping[col_type] = (
                        default_idx
                        if isinstance(default_idx, int) and default_idx < max_cols
                        else max_cols - 1
                    )
                else:
                    mapping[col_type] = default_idx
        self._has_header = has_header
        self._header_row_index = best_index if has_header else None
        self.last_detection = {
            'header_found': has_header,
            'header_row': best_index if has_header else None,
            'matches': max(0, best_score),
            'confidence': max(0.0, min(1.0, best_score / 5.0)),
        }

        return mapping

    def _detect_header_mapping(self, row: List[str]) -> Dict[str, int]:
        aliases = self.detection_config.get('aliases', {})
        priority = self.detection_config.get(
            'field_priority', list(COLUMN_TYPES)
        )
        mapping: Dict[str, int] = {}
        used_columns = set()
        for field in priority:
            for col_idx, header in enumerate(row):
                if col_idx in used_columns:
                    continue
                if self._header_matches(str(header), aliases.get(field, [])):
                    mapping[field] = col_idx
                    used_columns.add(col_idx)
                    break
        return mapping

    @staticmethod
    def _header_matches(header: str, aliases: List[str]) -> bool:
        normalized = ' '.join(header.casefold().split())
        for alias in aliases:
            value = str(alias or '').strip()
            if not value:
                continue
            if value.startswith('re:'):
                try:
                    if re.search(value[3:], normalized, re.IGNORECASE):
                        return True
                except re.error:
                    continue
            elif ' '.join(value.casefold().split()) in normalized:
                return True
        return False

    def get_available_columns(self, rows: List[List[str]]) -> List[int]:
        """Return available columns."""
        if not rows:
            return []

        max_cols = max(len(row) for row in rows)
        return list(range(max_cols))

    def parse_with_mapping(
        self,
        rows: List[List[str]],
        column_mapping: Dict[str, Optional[int]]
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Parse with mapping."""
        from collections import defaultdict

        char_data: Dict[str, Dict[str, Any]] = defaultdict(
            lambda: {"lines": 0, "raw": []}
        )
        lines_list = []

        # Detect whether a header is present and skip the first row only then
        start_idx = self._data_start_index(rows)
        data_rows = rows[start_idx:] if rows else []

        for row_idx, row in enumerate(data_rows):
            try:
                # Extract row data
                char_name = self._get_cell_value(row, column_mapping.get('character'))
                time_start = self._get_cell_value(row, column_mapping.get('time_start'))
                time_end = self._get_cell_value(row, column_mapping.get('time_end'))
                time_split = self._get_cell_value(row, column_mapping.get('time_split'))
                text = self._get_cell_value(row, column_mapping.get('text'))

                # Skip empty rows
                if not text and not char_name:
                    continue

                # Parse timings
                start_seconds = None
                end_seconds = None

                # Try a combined timing column first
                if time_split:
                    start_seconds, end_seconds = self._parse_split_time(time_split)

                # If that fails or split timing is absent, try separate timing columns
                if start_seconds is None:
                    start_seconds = self._parse_time(time_start)
                if end_seconds is None:
                    end_seconds = self._parse_time(time_end)

                # Use zero when no timings are available
                if start_seconds is None:
                    start_seconds = 0.0
                if end_seconds is None:
                    end_seconds = start_seconds + float(
                        self.detection_config.get('default_duration', 1.0)
                    )

                line_data = {
                    's': start_seconds,
                    'e': end_seconds,
                    'char': char_name or '',
                    'text': text or '',
                    's_raw': time_start or time_split or '',
                    'e_raw': time_end or '',
                }
                lines_list.append(line_data)

                # Update statistics
                if char_name:
                    char_data[char_name]["lines"] += 1
                    char_data[char_name]["raw"].append(line_data)

            except Exception as e:
                logger.warning(f"Error parsing row {row_idx}: {e}")
                continue

        # Calculate statistics
        # Do not merge DOCX rows; each row is a separate take
        # (the document is already merged; it is not subtitle timing)
        stats = []
        for char, info in char_data.items():
            char_lines = info["raw"]
            
            # For DOCX, rings equal lines because each replica is already merged
            rings = info["lines"]
            words = 0

            if char_lines:
                for line_data in char_lines:
                    words += len(line_data['text'].split())

            stats.append({
                "name": char,
                "lines": info["lines"],
                "rings": rings,
                "words": words
            })

        return stats, lines_list

    def _get_cell_value(self, row: List[str], col_idx: Optional[int]) -> Optional[str]:
        """Return cell value."""
        if col_idx is None or col_idx < 0 or col_idx >= len(row):
            return None
        return row[col_idx] if row[col_idx] else None

    def _parse_time(self, time_str: Optional[str]) -> Optional[float]:
        """Parse time."""
        if not time_str:
            return None

        time_str = time_str.strip()

        # Try several formats
        formats = [
            r'(\d{1,2}):(\d{2}):(\d{2})[,.](\d+)',  # HH:MM:SS,mmm
            r'(\d{1,2}):(\d{2})[,.](\d+)',  # MM:SS,mmm
            r'(\d{1,2}):(\d{2}):(\d{2})',  # HH:MM:SS
            r'(\d{1,2}):(\d{2})',  # MM:SS
        ]

        for pattern in formats:
            match = re.match(pattern, time_str)
            if match:
                groups = match.groups()
                try:
                    if len(groups) == 4:  # HH:MM:SS,mmm
                        h, m, s, ms = groups
                        seconds = int(h) * 3600 + int(m) * 60 + float(f"{s}.{ms}")
                        return seconds
                    elif len(groups) == 3:
                        if ':' in time_str and time_str.count(':') == 2:
                            h, m, s = groups
                            seconds = int(h) * 3600 + int(m) * 60 + float(s)
                            return seconds
                        else:
                            m, s, ms = groups
                            seconds = int(m) * 60 + float(f"{s}.{ms}")
                            return seconds
                    elif len(groups) == 2:  # MM:SS
                        m, s = groups
                        seconds = int(m) * 60 + float(s)
                        return seconds
                except (ValueError, IndexError):
                    continue

        # If no format matches, try the standard parser
        try:
            return srt_time_to_seconds(time_str.replace(',', '.'))
        except Exception:
            return None

    def _parse_split_time(self, time_str: Optional[str]) -> Tuple[Optional[float], Optional[float]]:
        """Parse split time."""
        if not time_str:
            return None, None

        time_str = time_str.strip()

        # Try to find a separator
        for separator in self.time_separators:
            # Escape special regex characters
            if separator == '\\t':
                separator = '\t'
            escaped_sep = re.escape(separator)

            # Find a separator with optional surrounding spaces
            pattern = rf'^(.+?)\s*{escaped_sep}\s*(.+)$'
            match = re.match(pattern, time_str)

            if match:
                start_part = match.group(1).strip()
                end_part = match.group(2).strip()

                start_seconds = self._parse_time(start_part)
                end_seconds = self._parse_time(end_part)

                if start_seconds is not None and end_seconds is not None:
                    return start_seconds, end_seconds

        return None, None

    def get_preview_data(
        self,
        rows: List[List[str]],
        column_mapping: Dict[str, Optional[int]],
        limit: int = 10
    ) -> List[Dict[str, Any]]:
        """Return preview data."""
        preview = []
        # Honor the header flag
        start_idx = self._data_start_index(rows)
        data_rows = rows[start_idx:start_idx + limit] if rows else []

        for row in data_rows:
            preview_row = {
                'raw': row,
                'mapped': {}
            }

            for col_type, col_idx in column_mapping.items():
                value = self._get_cell_value(row, col_idx)
                preview_row['mapped'][col_type] = value

                # Add the parsed value for timing fields
                if col_type in ['time_start', 'time_end']:
                    parsed = self._parse_time(value)
                    preview_row[f'{col_type}_parsed'] = parsed
                elif col_type == 'time_split':
                    start, end = self._parse_split_time(value)
                    preview_row['time_split_start_parsed'] = start
                    preview_row['time_split_end_parsed'] = end

            preview.append(preview_row)

        return preview

    def _data_start_index(self, rows: List[List[str]]) -> int:
        skipped = max(0, int(self.detection_config.get('rows_to_skip', 0)))
        header_index = getattr(self, '_header_row_index', None)
        if getattr(self, '_has_header', False):
            return min(
                len(rows), (header_index if header_index is not None else 0)
                + 1 + skipped
            )
        return min(len(rows), skipped)
