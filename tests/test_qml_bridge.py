"""Tests for the experimental QML bridge."""

from copy import deepcopy
from time import monotonic
import json
import os
from pathlib import Path
import tempfile

import pytest
from PySide6.QtCore import QCoreApplication, QSettings

from config.constants import MY_PALETTE
from services.audiobook_document_service import AudiobookDocumentService
from ui.qml_backend.app_bridge import AppBridge
from ui.qml_backend.features.teleprompter_bridge import (
    REAPER_ACTIVITY_TIMEOUT_SECONDS,
)
from ui.qml_backend.features.ui_state_bridge import UiStateBridge


def _app() -> QCoreApplication:
    app = QCoreApplication.instance()
    if app is None:
        app = QCoreApplication([])
    return app


@pytest.fixture(autouse=True)
def isolated_global_settings(monkeypatch, tmp_path):
    monkeypatch.setattr(
        "services.global_settings_service.SETTINGS_FILE",
        tmp_path / "global_settings.json",
    )
    settings_path = tmp_path / "ui-state.ini"

    def isolated_ui_state(parent=None):
        return UiStateBridge(
            QSettings(str(settings_path), QSettings.IniFormat), parent=parent
        )

    monkeypatch.setattr(
        "ui.qml_backend.app_bridge.UiStateBridge", isolated_ui_state
    )


def test_qml_bridge_starts_with_empty_project():
    _app()
    bridge = AppBridge()
    project = bridge.project

    assert project.name == "Новый проект"
    assert project.currentEpisode == ""
    assert project.episodesModel.rowCount() == 0
    assert bridge.project.episodesModel is project.episodesModel
    assert bridge.casting.actorsModel.rowCount() == 0
    assert bridge.casting.charactersModel.rowCount() == 0


def _configure_audiobook_project(bridge, tmp_path):
    source = str(tmp_path / "book.pdf")
    document = AudiobookDocumentService().create_document(source, [
        ("Глава 1", "<!DOCTYPE html><html><body><h1>Глава 1</h1><p>Первый текст.</p></body></html>"),
        ("Глава 2", "<!DOCTYPE html><html><body><h1>Глава 2</h1><p>Второй текст.</p></body></html>"),
    ])
    bridge._session.data.update({
        "project_kind": "audiobook",
        "episodes": {"Глава 1": source, "Глава 2": source},
        "audiobook_document": document,
        "actors": {"actor-1": {"name": "Актёр", "color": "#336699"}},
        "global_map": {},
        "audiobook_settings": {},
    })


def test_qml_audiobook_saves_html_markup_through_undo_command(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_audiobook_project(bridge, tmp_path)
    audiobook = bridge.audiobook
    audiobook.prepare()

    marked_html = (
        "<!DOCTYPE html><html><body><h1>Глава 1</h1><p>"
        '<span data-dm-character="Герой" data-dm-actor="actor-1">Первый</span> текст.'
        "</p></body></html>"
    )
    audiobook.updateEditorState(
        marked_html,
        '[{"character":"Автор","text":"Глава 1"},{"character":"Герой","text":"Первый"},{"character":"Автор","text":"текст."}]',
    )
    audiobook.setSlot(0, "Герой", "actor-1")

    assert audiobook.saveCurrent()
    document = bridge._session.data["audiobook_document"]
    lines = AudiobookDocumentService().lines(document, "Глава 1")
    assert any(line["char"] == "Герой" and line["text"] == "Первый" for line in lines)
    assert bridge._session.data["global_map"]["Герой"] == "actor-1"
    assert "book_chapters" not in bridge._session.data
    assert "audiobook_source" not in bridge._session.data
    assert bridge._session.data["episode_working_texts"] == {}

    bridge.project.undo()
    restored = AudiobookDocumentService().lines(
        bridge._session.data["audiobook_document"], "Глава 1"
    )
    assert not any(line["char"] == "Герой" for line in restored)

    bridge.project.redo()
    redone = AudiobookDocumentService().lines(
        bridge._session.data["audiobook_document"], "Глава 1"
    )
    assert any(line["char"] == "Герой" and line["text"] == "Первый" for line in redone)


def test_qml_audiobook_applies_reordered_chapter_structure(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_audiobook_project(bridge, tmp_path)
    audiobook = bridge.audiobook
    audiobook.prepare()
    audiobook.prepareChapterMarkup()
    audiobook.updateBoundaries(json.dumps({
        "sourceHtml": AudiobookDocumentService().combined_html(
            bridge._session.data["audiobook_document"]
        ),
        "chapters": [
            {"title": "Пролог", "html": "<!DOCTYPE html><html><body><h1>Пролог</h1><p>Второй текст.</p></body></html>"},
            {"title": "Глава 1", "html": "<!DOCTYPE html><html><body><h1>Глава 1</h1><p>Первый текст.</p></body></html>"},
        ],
    }, ensure_ascii=False))

    assert audiobook.applyChapterMarkup()
    assert AudiobookDocumentService().chapter_titles(
        bridge._session.data["audiobook_document"]
    ) == ["Пролог", "Глава 1"]
    assert set(bridge._session.data["episodes"]) == {"Пролог", "Глава 1"}

    bridge.project.undo()
    assert AudiobookDocumentService().chapter_titles(
        bridge._session.data["audiobook_document"]
    ) == ["Глава 1", "Глава 2"]
    assert set(bridge._session.data["episodes"]) == {"Глава 1", "Глава 2"}

    bridge.project.redo()
    assert AudiobookDocumentService().chapter_titles(
        bridge._session.data["audiobook_document"]
    ) == ["Пролог", "Глава 1"]


def test_qml_audiobook_teleprompter_edits_canonical_document(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_audiobook_project(bridge, tmp_path)
    bridge.project.selectEpisode("Глава 1")
    assert bridge.projectFiles.currentEpisodeSourceMissing is False
    bridge._session.project_service.current_project_path = str(
        tmp_path / "book.dub"
    )
    assert bridge.teleprompter.prepare("Глава 1")
    row = next(
        item for item in bridge.teleprompter.model.rows()
        if item["replicaText"] == "Первый текст."
    )

    assert bridge.teleprompter.editReplica(
        row["sourceIds"], "Герой", "Исправленный текст"
    )

    lines = AudiobookDocumentService().lines(
        bridge._session.data["audiobook_document"], "Глава 1"
    )
    assert any(
        line["char"] == "Герой" and line["text"] == "Исправленный текст"
        for line in lines
    )
    assert bridge._session.data["episode_working_texts"] == {}

    bridge.project.undo()
    restored = AudiobookDocumentService().lines(
        bridge._session.data["audiobook_document"], "Глава 1"
    )
    assert any(line["text"] == "Первый текст." for line in restored)

    bridge.project.redo()
    redone = AudiobookDocumentService().lines(
        bridge._session.data["audiobook_document"], "Глава 1"
    )
    assert any(
        line["char"] == "Герой" and line["text"] == "Исправленный текст"
        for line in redone
    )

    assert bridge.teleprompter.splitReplica(
        row["sourceIds"], "Исправленный", "текст", "Второй герой"
    )
    split_lines = AudiobookDocumentService().lines(
        bridge._session.data["audiobook_document"], "Глава 1"
    )
    assert any(line["char"] == "Второй герой" and line["text"] == "текст" for line in split_lines)

    bridge.project.undo()
    unsplit = AudiobookDocumentService().lines(
        bridge._session.data["audiobook_document"], "Глава 1"
    )
    assert not any(line["char"] == "Второй герой" for line in unsplit)

    bridge.project.redo()
    resplit = AudiobookDocumentService().lines(
        bridge._session.data["audiobook_document"], "Глава 1"
    )
    assert any(line["char"] == "Второй герой" for line in resplit)


def test_qml_audiobook_montage_edits_canonical_document(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_audiobook_project(bridge, tmp_path)
    bridge._session.project_service.current_project_path = str(
        tmp_path / "book.dub"
    )
    bridge.montage.prepare("Глава 1")
    line = next(
        item for item in bridge._script_text_service.load_episode_lines(
            bridge._session.data, "Глава 1"
        )
        if item["text"] == "Первый текст."
    )

    bridge.montage.updateText(line["working_id"], "Монтажная правка")

    updated = AudiobookDocumentService().lines(
        bridge._session.data["audiobook_document"], "Глава 1"
    )
    assert any(item["text"] == "Монтажная правка" for item in updated)
    assert bridge._session.data["episode_working_texts"] == {}

    bridge.project.undo()
    restored = AudiobookDocumentService().lines(
        bridge._session.data["audiobook_document"], "Глава 1"
    )
    assert any(item["text"] == "Первый текст." for item in restored)

    bridge.project.redo()
    redone = AudiobookDocumentService().lines(
        bridge._session.data["audiobook_document"], "Глава 1"
    )
    assert any(item["text"] == "Монтажная правка" for item in redone)


def test_qml_episode_model_uses_natural_and_audiobook_order(tmp_path):
    _app()
    bridge = AppBridge()
    source = str(tmp_path / "source.ass")
    bridge._session.data["episodes"] = {
        "10": source,
        "2": source,
        "1": source,
    }
    bridge.project.refresh_models()

    model = bridge.project.episodesModel
    assert [model.get(index)["name"] for index in range(3)] == ["1", "2", "10"]

    bridge._session.data.update({
        "project_kind": "audiobook",
        "episodes": {"Финал": source, "Пролог": source, "Глава 1": source},
        "audiobook_document": {
            "chapters": [
                {"title": title, "blocks": []}
                for title in ("Пролог", "Глава 1", "Финал")
            ]
        },
    })
    bridge.project.refresh_models()
    assert [model.get(index)["name"] for index in range(3)] == [
        "Пролог", "Глава 1", "Финал",
    ]


def _configure_teleprompter_project(bridge, tmp_path):
    bridge._session.data["episodes"] = {"1": str(tmp_path / "episode.ass")}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"},
        "actor-2": {"name": "Actor Two", "color": "#654321"},
    }
    bridge._session.data["global_map"] = {
        "Hero": "actor-1",
        "Villain": "actor-2",
    }
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "characters": {"Hero": {}, "Villain": {}},
            "lines": [
                {
                    "id": "line-1",
                    "start": 1.0,
                    "end": 2.0,
                    "character": "Hero",
                    "text": "First line",
                },
                {
                    "id": "line-2",
                    "start": 3.0,
                    "end": 4.0,
                    "character": "Villain",
                    "text": "Second line",
                },
            ],
        }
    }
    bridge.refresh()


def test_qml_bridge_prepares_and_navigates_teleprompter(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)

    prompter = bridge.teleprompter
    assert prompter.prepare("1")
    assert bridge.teleprompter.episode == "1"
    assert prompter.episode == "1"
    assert prompter.model is bridge.teleprompter.model
    assert bridge.teleprompter.model.rowCount() == 2
    assert bridge.teleprompter.model.rows()[0]["sourceIds"] == ["line-1"]
    assert bridge.teleprompter.model.rows()[0]["replicaText"] == "First line"

    prompter.navigate(1)
    assert bridge.teleprompter.time == 3.0
    assert prompter.time == 3.0
    assert bridge.teleprompter.currentIndex == 1

    prompter.setActorSelected("actor-2", False)
    assert [row["active"] for row in bridge.teleprompter.model.rows()] == [True, False]
    prompter.navigate(1)
    assert bridge.teleprompter.time == 1.0

    prompter.setActorSelected("actor-1", False)
    rows = bridge.teleprompter.model.rows()
    assert [row["active"] for row in rows] == [True, True]
    assert [row["colorActive"] for row in rows] == [False, False]


def test_qml_teleprompter_navigates_from_current_replica(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    bridge._session.data["episode_working_texts"]["1"]["lines"] = [
        {
            "id": "line-1",
            "start": 10.0,
            "end": 40.0,
            "character": "Hero",
            "text": "Long line",
        },
        {
            "id": "line-2",
            "start": 41.0,
            "end": 42.0,
            "character": "Villain",
            "text": "Next line",
        },
        {
            "id": "line-3",
            "start": 43.0,
            "end": 44.0,
            "character": "Hero",
            "text": "Following line",
        },
    ]
    prompter = bridge.teleprompter

    assert prompter.prepare("1")
    prompter._set_time(35.0, "reaper")
    assert prompter.currentIndex == 0

    prompter.navigate(1)
    assert prompter.time == 41.0

    prompter._set_time(35.0, "reaper")
    prompter.navigate(-1)
    assert prompter.time == 43.0

    prompter._set_time(41.5, "reaper")
    prompter.setActorSelected("actor-2", False)
    assert prompter.currentIndex == 1

    prompter.navigate(1)
    assert prompter.time == 43.0
    prompter._set_time(41.5, "reaper")
    prompter.navigate(-1)
    assert prompter.time == 10.0


def test_qml_teleprompter_defers_reaper_time_after_local_navigation(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    prompter = bridge.teleprompter

    assert prompter.prepare("1")
    prompter.jumpTo(3.0)
    assert prompter.positionOrigin == "local"
    prompter._on_osc_transport(True)

    prompter._pending_reaper_time = (3.0, float("inf"))
    prompter._on_osc_time(1.0)
    assert prompter.time == 3.0

    prompter._on_osc_time(3.0)
    assert prompter.time == 3.0
    assert prompter._pending_reaper_time is None

    prompter._on_osc_time(1.0)
    assert prompter.time == 1.0
    assert prompter.positionOrigin == "reaper"


def test_qml_teleprompter_follows_reaper_only_while_playing(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    prompter = bridge.teleprompter

    assert prompter.prepare("1")
    config = prompter.config
    config["sync_play_only"] = True
    bridge._global_settings_service.set_default_prompter_config(config)
    prompter._on_osc_time(3.0)
    assert prompter.time == 0.0

    prompter._on_osc_transport(True)
    assert prompter.time == 3.0

    prompter._on_osc_transport(False)
    prompter._on_osc_time(1.0)
    assert prompter.time == 3.0

    config = prompter.config
    config["sync_play_only"] = False
    bridge._global_settings_service.set_default_prompter_config(config)
    prompter._on_osc_time(1.0)
    assert prompter.time == 1.0



def test_qml_teleprompter_refresh_marks_position_as_internal(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    prompter = bridge.teleprompter

    assert prompter.prepare("1")
    prompter._on_osc_transport(True)
    prompter._on_osc_time(3.0)
    assert prompter.positionOrigin == "reaper"

    prompter.refresh()
    assert prompter.positionOrigin == "internal"


def test_qml_teleprompter_does_not_offset_reaper_input(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    prompter = bridge.teleprompter
    config = prompter.config
    config.update({
        "sync_in": True,
        "sync_out": True,
        "reaper_offset_enabled": True,
        "reaper_offset_seconds": -2.0,
    })
    bridge._global_settings_service.set_default_prompter_config(config)

    assert prompter.prepare("1")
    prompter._on_osc_transport(True)
    prompter._set_time(60.0, "local")
    prompter._pending_reaper_time = (58.0, float("inf"))

    prompter._on_osc_time(58.0)
    assert prompter.time == 60.0

    prompter._on_osc_time(58.1)
    assert prompter.time == 58.1
    assert prompter.positionOrigin == "reaper"


def test_qml_teleprompter_retries_osc_after_worker_error(tmp_path, monkeypatch):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    prompter = bridge.teleprompter
    config = prompter.config
    config["osc_enabled"] = True
    bridge._global_settings_service.set_default_prompter_config(config)
    prompter._episode = "1"
    worker = object()
    prompter._osc_worker = worker
    prompter._osc_client = object()
    prompter._osc_runtime_config = (8000, 9000)

    prompter._on_osc_error(worker, "port already in use")

    assert prompter.reaperConnectionState == "error"
    assert prompter.oscStatus == "Ошибка OSC: port already in use"
    assert prompter._osc_client is None
    assert prompter._osc_runtime_config is None

    starts = []
    monkeypatch.setattr(prompter, "_start_osc", lambda: starts.append(True))
    prompter.notify_global_config_changed()
    assert starts == [True]


def test_qml_teleprompter_restarts_osc_without_resetting_reader_state(
    tmp_path, monkeypatch
):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    prompter = bridge.teleprompter
    config = prompter.config
    config["osc_enabled"] = True
    bridge._global_settings_service.set_default_prompter_config(config)

    starts = []
    monkeypatch.setattr(prompter, "_start_osc", lambda: starts.append(True))
    assert prompter.prepare("1")
    prompter._set_time(3.0, "local")
    starts.clear()

    prompter.restartOsc()

    assert starts == [True]
    assert prompter.time == 3.0
    assert prompter.currentIndex == 1
    assert prompter.positionOrigin == "local"


def test_qml_teleprompter_restarts_osc_only_after_transport_changes(
    tmp_path, monkeypatch
):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    prompter = bridge.teleprompter
    prompter._episode = "1"

    config = prompter.config
    config["osc_enabled"] = True
    bridge._global_settings_service.set_default_prompter_config(config)
    prompter._osc_runtime_config = (8000, 9000)
    starts = []
    monkeypatch.setattr(prompter, "_start_osc", lambda: starts.append(True))

    prompter.notify_global_config_changed()
    assert starts == []

    config["port_out"] = 9001
    bridge._global_settings_service.set_default_prompter_config(config)
    prompter.notify_global_config_changed()
    assert starts == [True]


def test_qml_teleprompter_keeps_replicas_active_without_actors(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    bridge._session.data["actors"] = {}

    assert bridge.teleprompter.prepare("1")
    rows = bridge.teleprompter.model.rows()
    assert [row["active"] for row in rows] == [True, True]
    assert [row["colorActive"] for row in rows] == [False, False]


def test_qml_bridge_teleprompter_edits_and_splits_with_undo(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    bridge._session.project_service.current_project_path = str(
        tmp_path / "prompter-edit.dub"
    )
    bridge.teleprompter.prepare("1")
    bridge.reports.search("First line")
    bridge.reports.prepareSummary("")
    assert bridge.reports.searchResultCount == 1
    assert bridge.reports.summaryModel.rows()[0]["actor"] == "Actor One"

    assert bridge.teleprompter.editReplica(["line-1"], "Narrator", "Changed")
    edited = bridge._session.data["episode_working_texts"]["1"]["lines"][0]
    assert edited["display_character"] == "Narrator"
    assert edited["text"] == "Changed"
    assert bridge.reports.searchResultCount == 0
    assert bridge.reports.summaryModel.rows()[-1]["unassigned"] is True
    backups = list((tmp_path / ".backups").glob(
        "prompter-edit_editing_episode_1_*.dub_backup"
    ))
    assert len(backups) == 1

    before = deepcopy(
        bridge._session.data["episode_working_texts"]["1"]["lines"]
    )
    errors = []
    bridge.teleprompter.errorRequested.connect(errors.append)
    assert not bridge.teleprompter.editReplica(
        ["line-1", "line-2"], "Narrator", "Ambiguous replacement"
    )
    assert bridge._session.data["episode_working_texts"]["1"]["lines"] == before
    assert "объединённой реплики" in errors[-1]

    bridge.project.undo()
    restored = bridge._session.data["episode_working_texts"]["1"]["lines"][0]
    assert "display_character" not in restored
    assert restored["text"] == "First line"

    bridge.project.redo()
    redone = bridge._session.data["episode_working_texts"]["1"]["lines"][0]
    assert redone["display_character"] == "Narrator"
    assert redone["text"] == "Changed"
    bridge.project.undo()

    assert bridge.teleprompter.splitReplica(
        ["line-1"], "First", "line", "Narrator"
    )
    lines = bridge._session.data["episode_working_texts"]["1"]["lines"]
    assert len(lines) == 3
    assert lines[1]["display_character"] == "Narrator"
    assert lines[1]["text"] == "line"

    bridge.project.undo()
    assert len(bridge._session.data["episode_working_texts"]["1"]["lines"]) == 2
    bridge.project.redo()
    assert len(bridge._session.data["episode_working_texts"]["1"]["lines"]) == 3
    assert len(list((tmp_path / ".backups").glob(
        "prompter-edit_editing_episode_1_*.dub_backup"
    ))) == 1


def test_qml_multiple_actors_render_and_undo_across_casting_tools(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    bridge._session.current_episode = "1"

    bridge.casting.addActorToCharacter("Hero", "actor-2")

    assert bridge._session.data["global_map"]["Hero"] == [
        "actor-1", "actor-2",
    ]
    character = bridge.casting.charactersModel.rows()[0]
    assert [entry["name"] for entry in character["actorEntries"]] == [
        "Actor One", "Actor Two",
    ]

    assert bridge.teleprompter.prepare("1")
    hero = bridge.teleprompter.model.rows()[0]
    assert hero["active"] is True
    assert hero["colorActive"] is False

    bridge.teleprompter.setActorSelected("actor-2", False)
    hero = bridge.teleprompter.model.rows()[0]
    assert hero["active"] is True
    assert hero["colorActive"] is True
    assert hero["actorColor"] == "#123456"

    bridge.montage.prepare("1")
    assert bridge.montage.model.rows()[0]["background"] == "transparent"
    bridge.montage.setActorHighlighted("actor-2", False)
    assert bridge.montage.model.rows()[0]["background"] != "transparent"

    bridge.project.undo()
    assert bridge._session.data["global_map"]["Hero"] == "actor-1"


def test_qml_remove_one_actor_from_multi_cast_is_undoable(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    bridge._session.current_episode = "1"
    bridge.casting.addActorToCharacter("Hero", "actor-2")

    bridge.casting.removeActorFromCharacter("Hero", "actor-1")

    assert bridge._session.data["global_map"]["Hero"] == "actor-2"
    character = bridge.casting.charactersModel.rows()[0]
    assert [entry["id"] for entry in character["actorEntries"]] == ["actor-2"]

    bridge.project.undo()
    assert bridge._session.data["global_map"]["Hero"] == [
        "actor-1", "actor-2",
    ]


def test_qml_deleting_one_multi_cast_actor_keeps_the_other(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    bridge._session.current_episode = "1"
    bridge.casting.addActorToCharacter("Hero", "actor-2")

    bridge.casting.deleteActor("actor-1")

    assert bridge._session.data["global_map"]["Hero"] == "actor-2"
    assert "actor-1" not in bridge._session.data["actors"]

    bridge.project.undo()
    assert bridge._session.data["global_map"]["Hero"] == [
        "actor-1", "actor-2",
    ]


def test_qml_deleting_multiple_actors_is_one_undoable_operation():
    _app()
    bridge = AppBridge()
    bridge._session.data.update({
        "actors": {
            "actor-1": {"name": "One", "color": "#111111"},
            "actor-2": {"name": "Two", "color": "#222222"},
        },
        "global_map": {
            "Hero": "actor-1",
            "Pair": ["actor-1", "actor-2"],
        },
        "episode_actor_map": {"1": {"Guest": "actor-2"}},
    })
    bridge.casting.refresh()

    assert bridge.casting.actorRoleAssignments(["actor-1", "actor-2"]) == [
        {"actorName": "One", "rolesText": "Hero, Pair"},
        {"actorName": "Two", "rolesText": "Guest, Pair"},
    ]

    bridge.casting.deleteActors(["actor-1", "actor-2"])

    assert bridge._session.data["actors"] == {}
    assert bridge._session.data["global_map"] == {}
    assert bridge._session.data["episode_actor_map"] == {"1": {}}

    bridge.project.undo()
    assert set(bridge._session.data["actors"]) == {"actor-1", "actor-2"}
    assert bridge._session.data["global_map"] == {
        "Hero": "actor-1",
        "Pair": ["actor-1", "actor-2"],
    }
    assert bridge._session.data["episode_actor_map"] == {
        "1": {"Guest": "actor-2"}
    }


def test_qml_bridge_teleprompter_settings_and_presets(tmp_path, monkeypatch):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    bridge.teleprompter.prepare("1")

    refresh_calls = []
    original_refresh = bridge.refresh
    monkeypatch.setattr(bridge, "refresh", lambda: refresh_calls.append(True))
    bridge.teleprompter.setConfigValue("f_text", 52)
    bridge.teleprompter.setConfigValue("colors.bg", "#102030")
    assert refresh_calls == []
    monkeypatch.setattr(bridge, "refresh", original_refresh)
    assert bridge.teleprompter.config["f_text"] == 52
    assert bridge.teleprompter.config["colors"]["bg"] == "#102030"
    global_prompter = bridge._global_settings_service.get_default_prompter_config()
    assert global_prompter["layout_font_sizes"]["Сценарий 1"]["f_text"] == 52
    assert global_prompter["colors"]["bg"] == "#102030"

    bridge.teleprompter.setConfigValue("layout_type", "Сценарий 2")
    assert bridge.teleprompter.config["f_text"] == 36
    assert bridge.teleprompter.config["bold_char"] is True
    assert bridge.teleprompter.config["bold_text"] is False
    bridge.teleprompter.setConfigValue("f_text", 64)
    bridge.teleprompter.setConfigValue("bold_char", False)
    bridge.teleprompter.setConfigValue("bold_text", True)
    bridge.teleprompter.setConfigValue("layout_type", "Сценарий 1")
    assert bridge.teleprompter.config["f_text"] == 52
    assert bridge.teleprompter.config["bold_char"] is True
    assert bridge.teleprompter.config["bold_text"] is False
    profiles = bridge._global_settings_service.get_default_prompter_config()[
        "layout_font_sizes"
    ]
    assert profiles["Сценарий 2"]["f_text"] == 64
    bold_profiles = bridge._global_settings_service.get_default_prompter_config()[
        "layout_font_bold"
    ]
    assert bold_profiles["Сценарий 2"]["bold_char"] is False
    assert bold_profiles["Сценарий 2"]["bold_text"] is True

    bridge.teleprompter.setConfigValue("show_actor", False)
    bridge.teleprompter.setConfigValue("show_replica", False)
    bridge.teleprompter.setConfigValue("show_block_borders", True)
    bridge.teleprompter.setConfigValue("hide_leading_timecode_zeros", True)
    global_prompter = bridge._global_settings_service.get_default_prompter_config()
    assert global_prompter["show_actor"] is False
    assert global_prompter["show_replica"] is False
    assert global_prompter["show_block_borders"] is True
    assert global_prompter["hide_leading_timecode_zeros"] is True
    assert bridge.teleprompter.config["show_actor"] is False
    assert bridge.teleprompter.config["show_replica"] is False
    assert bridge.teleprompter.config["show_block_borders"] is True
    assert bridge.teleprompter.config["hide_leading_timecode_zeros"] is True

    bridge.teleprompter.savePreset(0)
    assert bridge.teleprompter.presetModel.rows()[0]["filled"]

    bridge.teleprompter.setConfigValue("colors.bg", "#000000")
    bridge.teleprompter.applyOrSavePreset(0)
    assert bridge.teleprompter.config["colors"]["bg"] == "#102030"
    assert bridge._global_settings_service.get_default_prompter_config()[
        "colors"
    ]["bg"] == "#102030"

    bridge.teleprompter.setConfigValue("colors.bg", "#405060")
    bridge.teleprompter.savePreset(0)
    assert bridge._global_settings_service.get_prompter_color_presets()[0][
        "bg"
    ] == "#405060"


def test_qml_teleprompter_uses_global_reaper_osc_settings(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    bridge._session.data["prompter_config"].update({
        "port_in": 7100,
        "port_out": 7101,
        "sync_in": False,
    })
    global_prompter = dict(bridge.settings.globalPrompterConfig)
    global_prompter.update({
        "port_in": 8100,
        "port_out": 8101,
        "sync_in": True,
        "osc_enabled": False,
    })

    assert bridge.settings.applyGlobalSettingsBundle(
        "ru",
        bridge.settings.audiobookKeywords,
        bridge.settings.globalMontageConfig,
        global_prompter,
    )

    assert bridge.teleprompter.config["port_in"] == 8100
    assert bridge.teleprompter.config["port_out"] == 8101
    assert bridge.teleprompter.config["sync_in"] is True
    bridge.teleprompter.setConfigValue("port_in", 9100)
    assert bridge.teleprompter.config["port_in"] == 9100
    assert bridge._session.data["prompter_config"]["port_in"] == 7100


def test_qml_teleprompter_sync_toggles_update_global_settings(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    bridge._session.data["prompter_config"].update({
        "sync_in": False,
        "sync_out": False,
    })

    assert bridge.settings.setPrompterSyncEnabled("sync_in", True)
    assert bridge.settings.setPrompterSyncEnabled("sync_out", True)
    assert bridge.settings.setPrompterSyncEnabled("sync_play_only", True)
    assert bridge.settings.setPrompterPageScrollMode(True)

    saved = bridge._global_settings_service.load_settings()
    assert saved["default_prompter_config"]["sync_in"] is True
    assert saved["default_prompter_config"]["sync_out"] is True
    assert saved["default_prompter_config"]["sync_play_only"] is True
    assert saved["default_prompter_config"]["page_scroll_mode"] is True
    assert bridge.teleprompter.config["sync_in"] is True
    assert bridge.teleprompter.config["sync_out"] is True
    assert bridge.teleprompter.config["page_scroll_mode"] is True
    assert bridge._session.data["prompter_config"]["sync_in"] is False
    assert bridge._session.data["prompter_config"]["page_scroll_mode"] is False
    assert not bridge.settings.setPrompterSyncEnabled("port_in", True)


def test_qml_teleprompter_saves_page_gap_prefetch_threshold(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    prompter = bridge.teleprompter

    prompter.setConfigValue("page_gap_prefetch_seconds", 2.5)
    assert prompter.config["page_gap_prefetch_seconds"] == 2.5

    prompter.setConfigValue("page_gap_prefetch_seconds", 99.0)
    assert prompter.config["page_gap_prefetch_seconds"] == 60.0

    prompter.setConfigValue("page_gap_prefetch_delay_seconds", 1.5)
    assert prompter.config["page_gap_prefetch_delay_seconds"] == 1.5

    prompter.setConfigValue("page_gap_prefetch_delay_seconds", 99.0)
    assert prompter.config["page_gap_prefetch_delay_seconds"] == 60.0

    prompter.setConfigValue("page_target_highlight_enabled", False)
    assert prompter.config["page_target_highlight_enabled"] is False

    prompter.setConfigValue("colors.page_target_highlight", "#336699")
    assert prompter.config["colors"]["page_target_highlight"] == "#336699"


def test_qml_teleprompter_reaper_indicator_tracks_osc_activity(tmp_path):
    _app()
    bridge = AppBridge()
    _configure_teleprompter_project(bridge, tmp_path)
    prompter = bridge.teleprompter
    global_config = bridge._global_settings_service.get_default_prompter_config()
    global_config["osc_enabled"] = True
    bridge._global_settings_service.set_default_prompter_config(global_config)
    prompter._osc_worker = object()

    prompter._refresh_reaper_connection_state()
    assert prompter.reaperConnectionState == "waiting"

    prompter._last_osc_activity = monotonic()
    prompter._refresh_reaper_connection_state()
    assert prompter.reaperConnectionState == "active"
    assert prompter.reaperConnectionText == "REAPER: синхронизация активна"

    prompter._last_osc_activity = (
        monotonic() - REAPER_ACTIVITY_TIMEOUT_SECONDS - 0.1
    )
    prompter._refresh_reaper_connection_state()
    assert prompter.reaperConnectionState == "lost"


def test_qml_bridge_normalizes_legacy_scenario_layout():
    _app()
    bridge = AppBridge()
    config = bridge._global_settings_service.get_default_export_config()
    config["layout_type"] = "Сценарий"
    bridge._global_settings_service.set_default_export_config(config)

    assert bridge.montage.config["layout_type"] == "Сценарий 1"


def test_qml_bridge_refreshes_episode_actor_and_line_models(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["project_name"] = "Demo"
    bridge._session.data["episodes"] = {"1": str(tmp_path / "episode.ass")}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"}
    }
    bridge._session.data["global_map"] = {"Hero": "actor-1"}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {
                    "id": "line-1",
                    "start": 1.5,
                    "end": 3.0,
                    "character": "Hero",
                    "text": "Hello from QML",
                }
            ]
        }
    }

    bridge.refresh()

    assert bridge.project.currentEpisode == "1"
    assert bridge.project.episodesModel.rowCount() == 1
    assert bridge.casting.actorsModel.rowCount() == 1
    assert bridge.casting.linesModel.rowCount() == 1
    assert bridge.casting.charactersModel.rowCount() == 1
    assert bridge.casting.charactersModel.rows()[0]["character"] == "Hero"
    assert bridge.casting.charactersModel.rows()[0]["actor"] == "Actor One"
    assert bridge.casting.linesModel.rows()[0]["actor"] == "Actor One"
    assert bridge.casting.linesModel.rows()[0]["text"] == "Hello from QML"


def test_qml_bridge_filters_character_model(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": str(tmp_path / "episode.ass")}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"},
        "actor-2": {"name": "Actor Two", "color": "#654321"},
    }
    bridge._session.data["global_map"] = {
        "Hero": "actor-1",
        "Villain": "actor-2",
    }
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {"id": "line-1", "start": 1.0, "end": 2.0, "character": "Hero", "text": "Alpha words"},
                {"id": "line-2", "start": 2.0, "end": 3.0, "character": "Villain", "text": "Beta words"},
                {"id": "line-3", "start": 3.0, "end": 4.0, "character": "Narrator", "text": "Gamma words"},
            ]
        }
    }
    bridge.refresh()
    casting = bridge.casting

    assert [row["character"] for row in casting.charactersModel.rows()] == [
        "Hero",
        "Narrator",
        "Villain",
    ]
    assert [row["name"] for row in bridge.casting.actorFilterModel.rows()] == [
        "Все актёры",
        "Без актёра",
        "Actor One",
        "Actor Two",
    ]

    casting.setActorFilter("actor-1")
    assert [row["character"] for row in casting.charactersModel.rows()] == ["Hero"]

    casting.setActorFilter("")
    casting.setActorFilter("__unassigned__")
    assert [row["character"] for row in casting.charactersModel.rows()] == ["Narrator"]

    casting.setShowUnassignedOnly(False)
    casting.setSearchText("vill")
    assert [row["character"] for row in casting.charactersModel.rows()] == ["Villain"]


def test_qml_bridge_selected_character_stats_reset_when_filtered_out(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": str(tmp_path / "episode.ass")}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"},
    }
    bridge._session.data["global_map"] = {"Hero": "actor-1"}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {"id": "line-1", "start": 1.0, "end": 2.0, "character": "Hero", "text": "Alpha words"},
                {"id": "line-2", "start": 2.0, "end": 3.0, "character": "Narrator", "text": "Beta words"},
            ]
        }
    }
    bridge.refresh()

    bridge.casting.selectCharacter("Hero")

    assert bridge.casting.selectedCharacter == "Hero"
    assert "Actor One" in bridge.casting.selectedCharacterStats
    assert "Реплик: 1" in bridge.casting.selectedCharacterStats

    bridge.casting.setShowUnassignedOnly(True)

    assert bridge.casting.selectedCharacter == ""
    assert bridge.casting.selectedCharacterStats == "Выберите персонажа в таблице"


def test_qml_bridge_actor_commands_use_undo_stack():
    _app()
    bridge = AppBridge()
    casting = bridge.casting

    casting.addActorWithDetails("New Actor", "#ABCDEF", "F")

    actor_rows = casting.actorsModel.rows()
    assert len(actor_rows) == 1
    actor_id = actor_rows[0]["id"]
    assert actor_rows[0]["name"] == "New Actor"
    assert actor_rows[0]["color"] == "#ABCDEF"
    assert actor_rows[0]["gender"] == "Ж"
    assert bridge.project.canUndo
    assert not bridge.project.canRedo

    casting.renameActor(actor_id, "Renamed Actor")
    assert bridge.casting.actorsModel.rows()[0]["name"] == "Renamed Actor"

    casting.updateActorColor(actor_id, "#123456")
    assert bridge.casting.actorsModel.rows()[0]["color"] == "#123456"

    casting.updateActorGender(actor_id, "M")
    assert bridge.casting.actorsModel.rows()[0]["gender"] == "М"

    bridge.project.undo()
    assert bridge.casting.actorsModel.rows()[0]["gender"] == "Ж"

    bridge.project.redo()
    assert bridge.casting.actorsModel.rows()[0]["gender"] == "М"

    bridge.project.undo()
    bridge.project.undo()
    assert bridge.casting.actorsModel.rows()[0]["color"] == "#ABCDEF"
    bridge.project.undo()
    assert bridge.casting.actorsModel.rows()[0]["name"] == "New Actor"

    casting.deleteActor(actor_id)
    assert bridge.casting.actorsModel.rowCount() == 0

    bridge.project.undo()
    assert bridge.casting.actorsModel.rows()[0]["name"] == "New Actor"


def test_qml_casting_sorts_project_actors_by_columns():
    _app()
    bridge = AppBridge()
    bridge._session.data["actors"] = {
        "actor-b": {"name": "Beta", "gender": "М"},
        "actor-a": {"name": "Alpha", "gender": "Ж"},
    }
    bridge._session.data["global_map"] = {
        "Hero": "actor-b",
        "Narrator": "actor-b",
    }
    bridge.casting.refresh()

    assert [row["name"] for row in bridge.casting.actorsModel.rows()] == [
        "Alpha", "Beta",
    ]

    bridge.casting.setActorSort("roleCount")
    assert bridge.casting.actorSortKey == "roleCount"
    assert [row["roleCount"] for row in bridge.casting.actorsModel.rows()] == [
        0, 2,
    ]

    bridge.casting.setActorSort("roleCount")
    assert not bridge.casting.actorSortAscending
    assert [row["roleCount"] for row in bridge.casting.actorsModel.rows()] == [
        2, 0,
    ]


def test_qml_actor_library_sorts_global_actors_independently():
    _app()
    bridge = AppBridge()
    library = bridge.actorLibrary
    library.addGlobalActor("Beta", "М")
    library.addGlobalActor("Alpha", "Ж")

    assert [row["name"] for row in library.globalActorsModel.rows()] == [
        "Alpha", "Beta",
    ]

    library.setActorSort("gender")
    assert library.actorSortKey == "gender"
    assert [row["gender"] for row in library.globalActorsModel.rows()] == [
        "Ж", "М",
    ]

    library.setActorSort("gender")
    assert not library.actorSortAscending
    assert [row["gender"] for row in library.globalActorsModel.rows()] == [
        "М", "Ж",
    ]


def test_qml_table_sort_preferences_survive_a_new_bridge(tmp_path):
    _app()
    settings_path = tmp_path / "ui-state.ini"
    ui_state = UiStateBridge(QSettings(str(settings_path), QSettings.IniFormat))
    bridge = AppBridge(ui_state=ui_state)

    bridge.casting.setCharacterSort("words")
    bridge.casting.setCharacterSort("words")
    bridge.casting.setActorSort("gender")
    bridge.actorLibrary.setActorSort("status")
    bridge.actorLibrary.setActorSort("status")

    restored_state = UiStateBridge(
        QSettings(str(settings_path), QSettings.IniFormat)
    )
    restored = AppBridge(ui_state=restored_state)

    assert restored.casting.characterSortKey == "words"
    assert not restored.casting.characterSortAscending
    assert restored.casting.actorSortKey == "gender"
    assert restored.casting.actorSortAscending
    assert restored.actorLibrary.actorSortKey == "status"
    assert not restored.actorLibrary.actorSortAscending

def test_qml_bridge_project_name_uses_undo_stack():
    _app()
    bridge = AppBridge()

    project = bridge.project
    project.name = "QML Project"

    assert project.name == "QML Project"
    assert project.canUndo

    project.undo()
    assert project.name == "Новый проект"

    project.redo()
    assert project.name == "QML Project"


def test_qml_project_settings_are_atomic_and_undoable():
    _app()
    bridge = AppBridge()
    original_metadata = dict(bridge._session.data["metadata"])
    assert bridge.settings.applyProjectSettingsFull("Settings Project", "Author", "Studio")

    assert bridge.project.name == "Settings Project"
    assert bridge.settings.projectAuthor == "Author"
    assert bridge.settings.projectStudio == "Studio"
    assert bridge.project.canUndo

    bridge.project.undo()

    assert bridge.project.name == "Новый проект"
    assert bridge._session.data["metadata"] == original_metadata

    bridge.project.redo()
    assert bridge.project.name == "Settings Project"


def test_audiobook_temporary_documents_use_platform_temp_directory():
    _app()
    bridge = AppBridge()

    temporary, url = bridge.audiobook._temporary_document(
        "<html></html>", "dm-test-XXXXXX.html"
    )

    assert temporary.isOpen()
    assert os.path.samefile(Path(url.toLocalFile()).parent, tempfile.gettempdir())


def test_qml_project_settings_only_update_project_metadata():
    _app()
    bridge = AppBridge()
    original_export = dict(bridge._session.data["export_config"])
    original_prompter = dict(bridge._session.data["prompter_config"])
    assert bridge.settings.applyProjectSettingsFull("Bundle Project", "Author", "Studio")

    assert bridge._session.data["export_config"] == original_export
    assert bridge._session.data["prompter_config"] == original_prompter

    bridge.project.undo()

    assert bridge.project.name == "Новый проект"
    assert bridge._session.data["export_config"] == original_export
    assert bridge._session.data["prompter_config"] == original_prompter
    assert not bridge.project.canUndo


def test_qml_project_settings_do_not_include_import_configs():
    _app()
    bridge = AppBridge()
    assert bridge.settings.applyProjectSettingsFull("Import Project", "Author", "Studio")
    assert not {
        "replica_merge_config", "ass_import_config", "srt_import_config",
        "docx_import_config",
    } & bridge._session.data.keys()

    bridge.project.undo()

    assert bridge.project.name == "Новый проект"
    assert not bridge.project.canUndo


def test_qml_global_settings_keep_russian_and_normalize_keywords():
    _app()
    bridge = AppBridge()

    assert bridge.settings.applyGlobalSettings(
        "en",
        "Chapter\n Глава \nchapter\nPart, Раздел",
    )

    saved = bridge._global_settings_service.load_settings()
    assert saved["language"] == "ru"
    assert saved["audiobook_config"] == {
        "chapter_keywords": ["Chapter", "Глава", "Part", "Раздел"]
    }
    assert bridge.settings.audiobookKeywords == "Chapter\nГлава\nPart\nРаздел"


def test_qml_global_settings_bundle_and_project_transfer(tmp_path):
    _app()
    bridge = AppBridge()
    montage = dict(bridge.settings.globalMontageConfig)
    montage.update({"layout_type": "Сценарий 2", "table_width_actor": 11.5})
    prompter = dict(bridge.settings.globalPrompterConfig)
    prompter.update({"f_text": 72, "show_header": True})

    assert bridge.settings.applyGlobalSettingsBundle(
        "ru", "Глава\nChapter", montage, prompter
    )
    saved = bridge._global_settings_service.load_settings()
    assert saved["default_export_config"]["layout_type"] == "Сценарий 2"
    assert saved["default_export_config"]["table_width_actor"] == 11.5
    assert saved["default_prompter_config"]["f_text"] == 72
    assert saved["default_prompter_config"]["show_header"] is True

    assert bridge._global_settings_service.get_default_export_config()["layout_type"] == "Сценарий 2"


def test_qml_global_settings_full_persists_unified_import_defaults():
    _app()
    bridge = AppBridge()
    merge = dict(bridge.settings.globalMergeConfig)
    merge.update({"fps": 30.0, "merge_gap": 60})
    ass = dict(bridge.settings.globalAssImportConfig)
    ass["character_separator"] = "/"
    srt = dict(bridge.settings.globalSrtImportConfig)
    srt["default_character"] = "Voice"
    docx = dict(bridge.settings.globalDocxImportConfig)
    docx["minimum_header_matches"] = 3

    assert bridge.settings.applyGlobalSettingsFull(
        "ru",
        "Глава\nChapter",
        bridge.settings.globalMontageConfig,
        bridge.settings.globalPrompterConfig,
        merge,
        ass,
        srt,
        docx,
    )

    saved = bridge._global_settings_service.load_settings()
    assert saved["default_replica_merge_config"]["fps"] == 30.0
    assert saved["default_replica_merge_config"]["merge_gap"] == 60.0
    assert saved["ass_import_config"]["character_separator"] == "/"
    assert saved["srt_import_config"]["default_character"] == "Voice"
    assert saved["docx_import_config"]["minimum_header_matches"] == 3


def test_qml_settings_manage_named_docx_import_presets():
    _app()
    bridge = AppBridge()
    config = dict(bridge.settings.globalDocxImportConfig)
    config["field_priority"] = [
        "text", "character", "time_split", "time_start", "time_end"
    ]
    config["fallback_mapping"] = {
        "character": 1,
        "time_start": None,
        "time_end": None,
        "time_split": 2,
        "text": 4,
    }

    assert bridge.settings.saveDocxImportPreset("Studio", config)
    presets = bridge.settings.globalDocxImportPresets
    assert presets[0]["name"] == "Studio"
    assert presets[0]["config"]["field_priority"][0] == "text"
    assert presets[0]["config"]["fallback_mapping"]["text"] == 4

    assert bridge.settings.deleteDocxImportPreset("Studio")
    assert bridge.settings.globalDocxImportPresets == []


def test_qml_global_backup_settings_reconfigure_project_service(tmp_path):
    _app()
    bridge = AppBridge()
    backup_config = {
        "enabled": True,
        "path_mode": "absolute",
        "directory": str(tmp_path / "central"),
        "interval_minutes": 12,
        "max_backups": 7,
    }

    assert bridge.settings.applyGlobalSettingsComplete(
        "ru",
        "Глава\nChapter",
        bridge.settings.globalMontageConfig,
        bridge.settings.globalPrompterConfig,
        bridge.settings.globalMergeConfig,
        bridge.settings.globalAssImportConfig,
        bridge.settings.globalSrtImportConfig,
        bridge.settings.globalDocxImportConfig,
        backup_config,
    )

    assert bridge.settings.globalBackupConfig == backup_config
    assert bridge._project_service.get_backup_config() == backup_config
    assert bridge.project._autosave_timer.interval() == 12 * 60_000


def test_qml_opens_backup_as_unsaved_full_project(tmp_path):
    _app()
    source = AppBridge()
    project_path = tmp_path / "source.dub"
    source.project.saveAs(str(project_path))
    source.project.name = "Полная резервная версия"
    source.casting.addActor("Актёр из копии")
    source.project.autoSave()
    backup_path = next(
        (tmp_path / ".backups").glob("source_*.dub_backup")
    )

    restored = AppBridge()
    restored.project.open(str(backup_path))

    assert restored.project.name == "Полная резервная версия"
    assert restored.casting.actorsModel.rowCount() == 1
    assert restored.project.path == ""
    assert restored.project.dirty
    assert restored.project.recentProjectsModel.rows()[1]["path"] != str(
        backup_path
    )
    assert "Сохраните её как обычный проект" in restored.statusText


def test_qml_open_embeds_external_working_text_and_rebases_paths(tmp_path):
    _app()
    (tmp_path / "Episode_01.ass").write_text(
        "[Events]\n"
        "Dialogue: 0,0:00:01.00,0:00:02.00,Default,Hero,0,0,0,,"
        "Source line\n",
        encoding="utf-8",
    )
    texts_dir = tmp_path / "Texts"
    texts_dir.mkdir()
    (texts_dir / "episode_1.json").write_text(json.dumps({
        "episode": "1",
        "lines": [{
            "id": "line-1",
            "start": 1.0,
            "end": 2.0,
            "character": "Hero",
            "text": "Portable line",
        }],
    }), encoding="utf-8")

    bridge = AppBridge()
    project_data = bridge._project_service.create_new_project("Portable")
    project_data.update({
        "project_name": "Portable",
        "project_folder": "/old/computer/Portable",
        "episodes": {"1": "Episode_01.ass"},
        "episode_texts": {"1": "Texts/episode_1.json"},
    })
    project_path = tmp_path / "portable.dub"
    project_path.write_text(
        json.dumps(project_data, ensure_ascii=False), encoding="utf-8"
    )

    bridge.project.open(str(project_path))

    assert bridge._session.data["project_folder"] == str(tmp_path.resolve())
    assert bridge._session.data["episode_working_texts"]["1"]["lines"][0][
        "text"
    ] == "Portable line"
    assert "1" not in bridge._session.data["episode_texts"]
    assert bridge.project.dirty is True
    working_row = next(
        row for row in bridge.projectFiles.filesModel.rows()
        if row["episode"] == "1" and row["kind"] == "working"
    )
    assert working_row["status"] == "Нет построчных реплик"

    migrated_path = tmp_path / "portable.dub"
    bridge.project.saveAs(str(migrated_path))

    saved = json.loads(migrated_path.read_text(encoding="utf-8"))
    payload = saved["episode_working_texts"]["1"]
    assert payload["lines"][0]["text"] == "Portable line"
    assert payload["source_lines_origin"] == "imported"
    assert payload["source_lines"][0]["text"] == "Source line"
    working_row = next(
        row for row in bridge.projectFiles.filesModel.rows()
        if row["episode"] == "1" and row["kind"] == "working"
    )
    assert working_row["status"] == "В проекте"


def test_qml_project_files_creates_missing_source_lines_with_undo(tmp_path):
    _app()
    source = tmp_path / "Episode_01.ass"
    source.write_text(
        "[Events]\n"
        "Dialogue: 0,0:00:01.00,0:00:02.00,Default,Hero,0,0,0,,"
        "Source line\n",
        encoding="utf-8",
    )
    bridge = AppBridge()
    bridge._session.data.update({
        "episodes": {"1": str(source)},
        "episode_working_texts": {
            "1": {
                "lines": [{
                    "id": "line-1",
                    "start": 1.0,
                    "end": 2.0,
                    "character": "Hero",
                    "text": "Edited working line",
                }],
            },
        },
    })
    bridge.refresh()

    bridge.projectFiles.createMissingSourceLines()

    payload = bridge._session.data["episode_working_texts"]["1"]
    assert payload["lines"][0]["text"] == "Edited working line"
    assert payload["source_lines"][0]["text"] == "Source line"
    assert payload["source_lines_origin"] == "imported"

    bridge.project.undo()

    assert "source_lines" not in bridge._session.data["episode_working_texts"]["1"]


def test_qml_global_import_profile_does_not_modify_project():
    _app()
    bridge = AppBridge()
    merge = dict(bridge.settings.globalMergeConfig)
    merge["merge"] = False
    ass = dict(bridge.settings.globalAssImportConfig)
    ass["strip_override_tags"] = False
    srt = dict(bridge.settings.globalSrtImportConfig)
    srt["keep_multiline"] = False
    docx = dict(bridge.settings.globalDocxImportConfig)
    docx["rows_to_skip"] = 2

    assert bridge.settings.saveImportConfigAsDefault(merge, ass, srt, docx)
    assert bridge._global_settings_service.get_replica_merge_config()["merge"] is False
    assert bridge._global_settings_service.get_ass_import_config()["strip_override_tags"] is False
    assert bridge._global_settings_service.get_srt_import_config()["keep_multiline"] is False
    assert bridge._global_settings_service.get_docx_import_config()["rows_to_skip"] == 2
    assert not bridge.project.canUndo


def test_qml_saves_current_settings_draft_as_global_default():
    _app()
    bridge = AppBridge()
    draft = dict(bridge.settings.projectPrompterConfig)
    draft["f_text"] = 81

    assert bridge.settings.saveConfigAsDefault("prompter", draft)

    assert bridge._global_settings_service.get_default_prompter_config()[
        "f_text"
    ] == 81


def test_qml_global_actor_base_export_import_roundtrip(tmp_path):
    _app()
    bridge = AppBridge()
    library = bridge.actorLibrary
    path = tmp_path / "global_actors.json"

    library.addGlobalActor("Actor One", "Ж")
    actor_id = library.globalActorsModel.rows()[0]["id"]
    assert library.exportGlobalActorBase(str(path))

    library.deleteGlobalActor(actor_id)
    assert library.globalActorsModel.rowCount() == 0

    assert library.importGlobalActorBase(str(path))
    assert library.globalActorsModel.rows()[0]["name"] == "Actor One"
    assert library.globalActorsModel.rows()[0]["gender"] == "Ж"

    library.addGlobalActor("Actor Two", "М")
    ids = [row["id"] for row in library.globalActorsModel.rows()]
    library.deleteGlobalActors(ids)
    assert library.globalActorsModel.rowCount() == 0


def test_qml_assignment_transfer_import_is_one_undoable_command(tmp_path):
    _app()
    source = AppBridge()
    source._session.data.update({
        "actors": {
            "actor-1": {
                "name": "Actor One",
                "color": "#123456",
                "gender": "М",
            },
        },
        "global_map": {"Hero": "actor-1"},
        "episode_actor_map": {"1": {"Guest": "actor-1"}},
        "episodes": {"1": "episode.ass"},
    })
    path = tmp_path / "assignments.json"
    assert source.actorLibrary.exportProjectAssignments(str(path))

    target = AppBridge()
    target._session.data["episodes"] = {"1": "target.ass"}
    assert target.actorLibrary.importProjectAssignments(str(path))

    assert target._session.data["actors"]["actor-1"]["name"] == "Actor One"
    assert target._session.data["actors"]["actor-1"]["gender"] == "М"
    assert target._session.data["global_map"] == {"Hero": "actor-1"}
    assert target._session.data["episode_actor_map"] == {
        "1": {"Guest": "actor-1"}
    }
    assert target.project.canUndo

    target.project.undo()

    assert target._session.data["actors"] == {}
    assert target._session.data["global_map"] == {}
    assert target._session.data["episode_actor_map"] == {}


def test_qml_bridge_save_project_as_adds_dub_suffix(tmp_path):
    _app()
    bridge = AppBridge()
    target = tmp_path / "saved_project"

    project = bridge.project
    project.saveAs(str(target))

    saved_path = target.with_suffix(".dub")
    assert saved_path.exists()
    assert project.path == str(saved_path)


def test_qml_bridge_tracks_recent_projects(tmp_path):
    _app()
    bridge = AppBridge()
    first = tmp_path / "first.dub"
    second = tmp_path / "second.dub"

    project = bridge.project
    project.saveAs(str(first))
    project.saveAs(str(second))

    recent_rows = project.recentProjectsModel.rows()
    assert recent_rows[0]["display"] == "Недавние проекты"
    assert recent_rows[1]["path"] == str(second)
    assert recent_rows[2]["path"] == str(first)
    assert recent_rows[-1]["path"] == "__clear__"

    project.openRecent(str(first))

    assert project.path == str(first)
    assert project.recentProjectsModel.rows()[1]["path"] == str(first)

    project.clearRecent()
    recent_rows = project.recentProjectsModel.rows()
    assert [row["display"] for row in recent_rows] == [
        "Недавние проекты",
        "Нет недавних проектов",
    ]


def test_qml_bridge_imports_subtitle_file_with_undo(tmp_path):
    _app()
    bridge = AppBridge()
    srt_path = tmp_path / "Episode_02.srt"
    srt_path.write_text(
        "1\n00:00:01,000 --> 00:00:02,000\nHero: Hello from import\n",
        encoding="utf-8",
    )

    project = bridge.project
    project.importSubtitle(str(srt_path))

    assert project.currentEpisode == "02"
    assert project.episodesModel.rowCount() == 1
    assert bridge.casting.linesModel.rowCount() == 1
    assert bridge.casting.charactersModel.rows()[0]["character"] == "Hero"
    assert bridge._session.data["episodes"]["02"] == str(srt_path)
    assert bridge._session.data["episode_working_texts"]["02"]["lines"][0]["text"] == "Hello from import"
    assert project.dirty
    assert project.canUndo

    project.undo()

    assert project.currentEpisode == ""
    assert project.episodesModel.rowCount() == 0
    assert bridge.casting.linesModel.rowCount() == 0
    assert bridge.casting.charactersModel.rowCount() == 0

    project.redo()

    assert project.currentEpisode == "02"
    assert project.episodesModel.rowCount() == 1
    assert bridge.casting.linesModel.rowCount() == 1


def test_qml_subtitle_import_prepares_unique_editable_episode_names(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"02": "existing.srt"}
    first = tmp_path / "Episode_02.srt"
    second = tmp_path / "Alternative_02.ass"
    first.write_text("", encoding="utf-8")
    second.write_text("", encoding="utf-8")

    importer = bridge.subtitleImport
    assert importer.prepare([str(first), str(second)])

    rows = importer.model.rows()
    assert [row["episode"] for row in rows] == ["02 2", "02 3"]
    assert importer.canImport

    importer.setEpisode(1, "02 2")
    assert not importer.canImport
    assert {row["status"] for row in importer.model.rows()} == {
        "Название повторяется"
    }

    importer.setEpisode(1, "03")
    assert importer.canImport


def test_qml_subtitle_import_is_atomic_and_undoable(tmp_path):
    _app()
    bridge = AppBridge()
    first = tmp_path / "Episode_01.srt"
    second = tmp_path / "Episode_02.srt"
    first.write_text(
        "1\n00:00:01,000 --> 00:00:02,000\nHero: First\n",
        encoding="utf-8",
    )
    second.write_text(
        "1\n00:00:03,000 --> 00:00:04,000\nGuest: Second\n",
        encoding="utf-8",
    )

    importer = bridge.subtitleImport
    assert importer.prepare([str(first), str(second)])
    assert importer.importAll()

    assert bridge.project.currentEpisode == "02"
    assert bridge._session.data["episodes"] == {
        "01": str(first),
        "02": str(second),
    }
    assert set(bridge._session.data["episode_working_texts"]) == {"01", "02"}
    assert bridge.project.canUndo

    bridge.project.undo()

    assert bridge._session.data["episodes"] == {}
    assert bridge._session.data.get("episode_working_texts", {}) == {}
    assert bridge.project.episodesModel.rowCount() == 0


def test_qml_subtitle_import_rejects_existing_episode_name(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": "existing.srt"}
    source = tmp_path / "Episode_2.srt"
    source.write_text("", encoding="utf-8")

    importer = bridge.subtitleImport
    assert importer.prepare([str(source)])
    importer.setEpisode(0, "1")

    assert not importer.canImport
    assert importer.model.rows()[0]["status"] == "Уже существует"
    assert not importer.importAll()


def _finish_converter(converter, limit=100):
    app = _app()
    for _ in range(limit):
        app.processEvents()
        if not converter.busy:
            return
    pytest.fail("Converter queue did not finish")


def test_qml_quick_converter_exports_standalone_files_without_project_mutation(
    tmp_path,
):
    _app()
    bridge = AppBridge()
    source = tmp_path / "quick_01.srt"
    source.write_text(
        "1\n00:00:01,000 --> 00:00:02,500\nHero: Quick line\n",
        encoding="utf-8",
    )
    unsupported = tmp_path / "notes.txt"
    unsupported.write_text("not subtitles", encoding="utf-8")
    original_project = deepcopy(bridge._session.data)

    converter = bridge.converter
    converter.setFormat("html", True)
    converter.setFormat("docx", False)
    converter.setFormat("pdf", False)
    assert converter.convert([str(source), str(unsupported)], False)
    _finish_converter(converter)

    output = tmp_path / "quick_01.html"
    assert output.exists()
    assert "Quick line" in output.read_text(encoding="utf-8")
    rows = converter.model.rows()
    assert rows[0]["statusKind"] == "success"
    assert rows[0]["outputPath"] == str(output)
    assert rows[1]["statusKind"] == "skipped"
    assert bridge._session.data == original_project

    assert converter.convert([str(source)], False)
    _finish_converter(converter)
    assert (tmp_path / "quick_01 (2).html").exists()


def test_qml_quick_converter_previews_then_converts(tmp_path):
    _app()
    bridge = AppBridge()
    source = tmp_path / "preview.srt"
    source.write_text(
        "1\n00:00:01,000 --> 00:00:02,500\nHero: Preview line\n",
        encoding="utf-8",
    )
    converter = bridge.converter
    converter.setFormat("html", True)
    converter.setFormat("docx", False)
    converter.setFormat("pdf", False)
    previews = []
    converter.previewRequested.connect(lambda: previews.append(True))

    assert converter.convert([str(source)], True)
    assert previews == [True]
    assert "Preview line" in converter.previewHtml
    assert converter.previewTitle == "preview.srt"
    assert not converter.busy

    converter.continueAfterPreview()
    _finish_converter(converter)
    assert (tmp_path / "preview.html").exists()


def test_qml_quick_converter_demo_settings_are_project_independent(
    tmp_path, monkeypatch,
):
    _app()
    settings_path = tmp_path / "global_settings.json"
    monkeypatch.setattr(
        "services.global_settings_service.SETTINGS_FILE", settings_path
    )
    bridge = AppBridge()
    converter = bridge.converter
    project_before = deepcopy(bridge._session.data)
    previews = []
    converter.previewRequested.connect(lambda: previews.append(True))

    assert converter.openSettingsPreview()
    assert converter.previewSettingsMode
    assert previews == [True]
    assert "Демонстрационная реплика" in converter.previewHtml

    converter.setPreviewOption("layout_type", "Сценарий 2")
    converter.setPreviewOption("f_text", 41)
    assert converter.savePreviewSettings()
    assert not converter.previewSettingsMode
    assert bridge._session.data == project_before

    saved = json.loads(settings_path.read_text(encoding="utf-8"))
    assert saved["quick_converter_config"]["layout_type"] == "Сценарий 2"
    assert saved["quick_converter_config"]["f_text"] == 41
    assert saved["default_export_config"] == project_before["export_config"]


def test_qml_quick_converter_applies_preview_settings_to_all_files(tmp_path):
    _app()
    bridge = AppBridge()
    paths = []
    for index in range(2):
        source = tmp_path / f"styled_{index}.srt"
        source.write_text(
            "1\n00:00:01,000 --> 00:00:02,500\nHero: Styled line\n",
            encoding="utf-8",
        )
        paths.append(str(source))

    converter = bridge.converter
    converter.setFormat("html", True)
    converter.setFormat("docx", False)
    converter.setFormat("pdf", False)
    assert converter.convert(paths, True)

    original_preview = converter.previewHtml
    converter.setPreviewOption("layout_type", "Сценарий 2")
    converter.setPreviewOption("f_text", 42)

    assert converter.previewConfig["layout_type"] == "Сценарий 2"
    assert converter.previewConfig["f_text"] == 42
    assert converter.previewHtml != original_preview

    converter.continueAfterPreview()
    _finish_converter(converter)

    for index in range(2):
        exported = tmp_path / f"styled_{index}.html"
        assert exported.exists()
        html = exported.read_text(encoding="utf-8")
        assert "Styled line" in html
        assert "42" in html


def test_qml_quick_converter_can_cancel_from_preview(tmp_path):
    _app()
    bridge = AppBridge()
    source = tmp_path / "cancel_preview.srt"
    source.write_text(
        "1\n00:00:01,000 --> 00:00:02,500\nHero: Cancelled line\n",
        encoding="utf-8",
    )
    converter = bridge.converter
    converter.setFormat("html", True)

    assert converter.convert([str(source)], True)
    converter.cancelPreview()

    assert not converter.busy
    assert converter.model.rows()[0]["status"] == "Отменено"
    assert not (tmp_path / "cancel_preview.html").exists()


def test_qml_quick_converter_cancels_before_next_file(tmp_path):
    _app()
    bridge = AppBridge()
    paths = []
    for index in range(2):
        path = tmp_path / f"cancel_{index}.srt"
        path.write_text(
            "1\n00:00:01,000 --> 00:00:02,500\nHero: Line\n",
            encoding="utf-8",
        )
        paths.append(str(path))

    converter = bridge.converter
    converter.setFormat("html", True)
    converter.setFormat("docx", False)
    converter.setFormat("pdf", False)
    assert converter.convert(paths, False)
    converter.cancel()
    _finish_converter(converter)

    assert [row["status"] for row in converter.model.rows()] == [
        "Отменено", "Отменено"
    ]
    assert not list(tmp_path.glob("cancel_*.html"))


def test_qml_bridge_imports_docx_with_preview_and_atomic_undo(tmp_path):
    from docx import Document

    _app()
    bridge = AppBridge()
    path = tmp_path / "Episode_03.docx"
    document = Document()
    table = document.add_table(rows=3, cols=3)
    for column, value in enumerate(("Персонаж", "Тайминг", "Текст")):
        table.cell(0, column).text = value
    for row, values in enumerate((
        ("Hero", "00:00:01,000 - 00:00:02,000", "First line"),
        ("Guest", "00:00:03,000 - 00:00:04,000", "Second line"),
    ), start=1):
        for column, value in enumerate(values):
            table.cell(row, column).text = value
    document.save(path)

    importer = bridge.docxImport
    assert importer.load(str(path))
    assert importer.suggestedEpisode == "03"
    assert importer.mapping["character"] == 0
    assert importer.mapping["time_split"] == 1
    assert importer.mapping["text"] == 2
    assert importer.previewModel.rowCount() == 2
    assert importer.importEpisode("03", False)

    assert bridge.project.currentEpisode == "03"
    assert bridge._session.data["episodes"]["03"] == str(path)
    assert bridge._session.data["episode_working_texts"]["03"]["lines"][0][
        "text"
    ] == "First line"
    assert bridge._global_settings_service.get_docx_import_config()["mapping"]["text"] == 2

    bridge.project.undo()

    assert "03" not in bridge._session.data["episodes"]
    assert "03" not in bridge._session.data["episode_working_texts"]


def test_qml_bridge_exposes_actor_palette():
    _app()
    bridge = AppBridge()

    assert bridge.casting.actorPalette == list(MY_PALETTE)


def test_qml_bridge_global_search_and_result_navigation(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {
        "1": str(tmp_path / "one.ass"),
        "2": str(tmp_path / "two.ass"),
    }
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {
                    "id": "line-1",
                    "start": 1.0,
                    "end": 2.0,
                    "character": "Hero",
                    "text": "First phrase",
                },
            ],
        },
        "2": {
            "lines": [
                {
                    "id": "line-2",
                    "start": 3.0,
                    "end": 4.0,
                    "character": "Narrator",
                    "text": "Hero returns here",
                },
            ],
        },
    }
    bridge.refresh()

    reports = bridge.reports
    count = reports.search("hero")

    assert count == 2
    assert reports.searchResultCount == 2
    assert [row["episode"] for row in reports.searchModel.rows()] == ["1", "2"]
    reports.setSearchSort("episode")
    assert reports.searchSortAscending is False
    assert [row["episode"] for row in reports.searchModel.rows()] == ["2", "1"]
    reports.setSearchSort("character")
    assert reports.searchSortKey == "character"
    assert [row["character"] for row in reports.searchModel.rows()] == [
        "Hero", "Narrator",
    ]

    bridge.casting.setSearchText("missing")
    reports.openResult("2", "Narrator")

    assert bridge.project.currentEpisode == "2"
    assert bridge.casting.searchText == ""
    assert bridge.casting.selectedCharacter == "Narrator"


def test_qml_bridge_prepares_and_exports_project_summary(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["project_name"] = "Demo"
    bridge._session.data["episodes"] = {"1": str(tmp_path / "one.ass")}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"},
    }
    bridge._session.data["global_map"] = {"Hero": "actor-1"}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {
                    "id": "line-1",
                    "start": 1.0,
                    "end": 2.0,
                    "character": "Hero",
                    "text": "Alpha beta",
                },
            ],
        },
    }
    bridge.refresh()

    reports = bridge.reports
    reports.prepareSummary("")

    assert reports.summaryTarget == ""
    assert reports.summaryModel.rows()[0]["actor"] == "Actor One"
    assert reports.summaryModel.rows()[0]["rings"] == 1
    assert reports.summaryModel.rows()[0]["words"] == 2
    assert reports.summaryModel.rows()[0]["roles"] == "Hero"
    reports.setSummarySort("words")
    assert reports.summarySortKey == "words"
    reports.setSummarySort("words")
    assert reports.summarySortAscending is False

    export_path = tmp_path / "project-summary"
    reports.exportProjectSummaryXlsx(str(export_path), "words")

    assert export_path.with_suffix(".xlsx").exists()
    assert reports.projectSummaryMetric == "words"


def test_qml_casting_moves_global_actor_to_project_with_undo():
    _app()
    bridge = AppBridge()
    casting = bridge.casting
    actor_library = bridge.actorLibrary

    actor_library.addGlobalActor("Global Voice", "Ж")
    global_row = actor_library.globalActorsModel.rows()[0]

    assert global_row["name"] == "Global Voice"
    assert global_row["gender"] == "Ж"
    assert global_row["color"] == "transparent"

    actor_library.addGlobalActorToProject(global_row["id"], "#4F81BD")

    project_actor = bridge._session.data["actors"][global_row["id"]]
    assert project_actor["name"] == "Global Voice"
    assert project_actor["gender"] == "Ж"
    assert project_actor["color"] == "#4F81BD"
    assert actor_library.globalActorsModel.rows()[0]["inProject"]

    bridge.project.undo()

    assert bridge._session.data["actors"] == {}
    assert not actor_library.globalActorsModel.rows()[0]["inProject"]


def test_qml_casting_moves_multiple_global_actors_to_project_with_undo():
    _app()
    bridge = AppBridge()
    actor_library = bridge.actorLibrary

    actor_library.addGlobalActor("Global Voice One", "Ж")
    actor_library.addGlobalActor("Global Voice Two", "М")
    rows = actor_library.globalActorsModel.rows()

    actor_library.addGlobalActorsToProject([row["id"] for row in rows])

    project_actors = bridge._session.data["actors"]
    assert len(project_actors) == 2
    colors = [actor["color"] for actor in project_actors.values()]
    assert all(color in MY_PALETTE for color in colors)
    assert len(set(colors)) == 2

    bridge.project.undo()

    assert bridge._session.data["actors"] == {}


def test_qml_actor_merge_replaces_every_project_reference_and_undoes():
    _app()
    bridge = AppBridge()
    library = bridge.actorLibrary
    bridge._session.data["actors"] = {
        "source": {"name": "Duplicate", "color": "#111111", "gender": "Ж"},
        "target": {"name": "Keeper", "color": "#222222", "gender": ""},
    }
    bridge._session.data["global_map"] = {"Hero": "source"}
    bridge._session.data["episode_actor_map"] = {
        "1": {"Guest": "source"},
    }
    bridge._session.data["export_config"]["highlight_ids_export"] = [
        "source", "target",
    ]
    bridge.refresh()

    assert library.mergeProjectActor("source", "project", "target")
    assert set(bridge._session.data["actors"]) == {"target"}
    assert bridge._session.data["actors"]["target"]["gender"] == "Ж"
    assert bridge._session.data["global_map"] == {"Hero": "target"}
    assert bridge._session.data["episode_actor_map"] == {
        "1": {"Guest": "target"},
    }
    assert bridge._session.data["export_config"]["highlight_ids_export"] == [
        "target",
    ]

    bridge.project.undo()
    assert set(bridge._session.data["actors"]) == {"source", "target"}
    assert bridge._session.data["global_map"] == {"Hero": "source"}


def test_qml_character_stats_cover_every_episode_and_sorting(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {
        "2": str(tmp_path / "two.ass"),
        "1": str(tmp_path / "one.ass"),
    }
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Яна", "color": "#123456"},
    }
    bridge._session.data["global_map"] = {"Hero": "actor-1"}
    bridge._session.data["episode_working_texts"] = {
        "1": {"lines": [
            {"character": "Hero", "text": "one two", "start": 1.0, "end": 3.0},
            {"character": "Other", "text": "one", "start": 4.0, "end": 5.0},
        ]},
        "2": {"lines": [
            {"character": "Hero", "text": "three four five", "start": 2.0, "end": 6.0},
        ]},
    }
    bridge.refresh()
    bridge.project.selectEpisode("1")
    casting = bridge.casting
    casting.selectCharacter("Hero")

    assert casting.characterEpisodeStatsModel.rows() == [
        {"episode": "1", "lines": 1, "rings": 1, "words": 2,
         "actor": "Яна", "scope": "Проект"},
        {"episode": "2", "lines": 1, "rings": 1, "words": 3,
         "actor": "Яна", "scope": "Проект"},
    ]
    assert "Реплик: 2" in casting.selectedCharacterStats
    assert casting.timelineDuration == 5.0
    assert casting.timelineActorsModel.rows() == [
        {"id": "actor-1", "name": "Яна", "actorColor": "#123456", "lane": 0},
        {"id": "", "name": "Без актёра", "actorColor": "#9A9A9A", "lane": 1},
    ]
    assert casting.timelineModel.rows() == [
        {
            "start": 1.0, "end": 3.0, "character": "Hero", "actor": "Яна",
            "actorId": "actor-1", "actorColor": "#123456", "lane": 0, "selected": True,
        },
        {
            "start": 4.0, "end": 5.0, "character": "Other", "actor": "Без актёра",
            "actorId": "", "actorColor": "#9A9A9A", "lane": 1, "selected": False,
        },
    ]
    casting.setTimelineSortMode("words")
    assert casting.timelineSortMode == "words"
    assert [row["name"] for row in casting.timelineActorsModel.rows()] == [
        "Яна", "Без актёра",
    ]
    casting.setTimelineSortMode("name")
    assert [row["name"] for row in casting.timelineActorsModel.rows()] == [
        "Без актёра", "Яна",
    ]
    casting.setTimelineSortMode("invalid")
    assert casting.timelineSortMode == "appearance"
    casting.setCharacterSort("words")
    assert [row["character"] for row in casting.charactersModel.rows()] == [
        "Other", "Hero",
    ]
    casting.setCharacterSort("words")
    assert [row["character"] for row in casting.charactersModel.rows()] == [
        "Hero", "Other",
    ]


def test_qml_casting_assigns_project_roles_atomically_with_undo(tmp_path):
    _app()
    bridge = AppBridge()
    roles = bridge.roles
    bridge._session.data["actors"] = {
        "actor-1": {"name": "One", "color": "#123456"},
        "actor-2": {"name": "Two", "color": "#654321"},
    }
    bridge._session.data["episodes"] = {"1": str(tmp_path / "one.ass")}
    bridge._session.data["global_map"] = {"Hero": "actor-1"}
    bridge._session.data["episode_actor_map"] = {
        "1": {"Hero": "actor-1", "Guest": "actor-1"},
    }
    bridge._session.data["episode_working_texts"] = {
        "1": {"lines": [
            {"character": "Hero", "text": "Hello"},
            {"character": "Guest", "text": "Hi"},
        ]},
    }
    bridge.refresh()

    roles.assign(["Hero", "Guest"], "actor-2")

    assert bridge._session.data["global_map"] == {
        "Hero": "actor-2",
        "Guest": "actor-2",
    }
    assert bridge._session.data["episode_actor_map"]["1"] == {}

    bridge.project.undo()

    assert bridge._session.data["global_map"] == {"Hero": "actor-1"}
    assert bridge._session.data["episode_actor_map"] == {
        "1": {"Hero": "actor-1", "Guest": "actor-1"},
    }


def test_qml_casting_syncs_global_actor_ids_and_preserves_project_color():
    _app()
    bridge = AppBridge()
    actor_library = bridge.actorLibrary
    actor_library.addGlobalActor("Shared Voice", "М")
    global_id = actor_library.globalActorsModel.rows()[0]["id"]
    bridge._session.data["actors"] = {
        "local-id": {
            "name": "Shared Voice", "color": "#ABCDEF", "gender": "",
        },
    }
    bridge._session.data["global_map"] = {"Hero": "local-id"}

    assert actor_library.syncProjectActorsWithGlobalBase() == 1
    assert bridge._session.data["actors"][global_id]["color"] == "#ABCDEF"
    assert bridge._session.data["actors"][global_id]["gender"] == "М"
    assert bridge._session.data["global_map"]["Hero"] == global_id

    bridge.project.undo()
    assert "local-id" in bridge._session.data["actors"]
    assert bridge._session.data["global_map"]["Hero"] == "local-id"


def test_qml_casting_builds_actor_role_stats_and_bulk_global_transfer(tmp_path):
    _app()
    bridge = AppBridge()
    roles = bridge.roles
    actor_library = bridge.actorLibrary
    bridge._session.data["actors"] = {
        "actor-1": {"name": "One", "color": "#123456", "gender": "Ж"},
        "actor-2": {"name": "Two", "color": "#654321", "gender": "М"},
    }
    bridge._session.data["episodes"] = {"1": str(tmp_path / "one.ass")}
    bridge._session.data["global_map"] = {"Hero": "actor-1"}
    bridge._session.data["episode_working_texts"] = {
        "1": {"lines": [
            {"character": "Hero", "text": "Alpha beta"},
            {"character": "Hero", "text": "Gamma"},
        ]},
    }
    bridge._session.data["replica_merge_config"] = {"merge": False}
    bridge.refresh()

    roles.prepareActorStats("actor-1")
    assert roles.actorStatsTitle == "One"
    assert roles.actorStatsModel.rows() == [
        {"name": "Hero", "rings": 2, "words": 3},
    ]

    actor_library.refreshProjectActorTransfer()
    assert len(actor_library.projectActorTransferModel.rows()) == 2
    actor_library.addProjectActorsToGlobal(["actor-1", "actor-2"])
    assert {
        actor["name"]
        for actor in bridge._global_settings_service.get_global_actor_base().values()
    } == {"One", "Two"}


def test_qml_bridge_previews_montage_and_saves_global_export_settings(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": str(tmp_path / "one.ass")}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"},
    }
    bridge._session.data["global_map"] = {"Hero": "actor-1"}
    bridge._session.data["replica_merge_config"] = {"merge": False}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {
                    "id": "line-1",
                    "start": 1.0,
                    "end": 2.0,
                    "character": "Hero",
                    "text": "Montage preview",
                },
            ],
        },
    }
    bridge.refresh()

    montage = bridge.montage
    montage.prepare("1")

    assert montage.episode == "1"
    assert montage.model.rowCount() == 1
    preview_row = montage.model.rows()[0]
    assert preview_row["character"] == "Hero"
    assert preview_row["actor"] == "Actor One"
    assert preview_row["text"] == "Montage preview"
    assert montage.count == 1
    assert "<table>" in montage.html
    assert "qrc:///qtwebchannel/qwebchannel.js" in montage.html
    assert "contenteditable='true'" in montage.html

    montage.setOption("layout_type", "Сценарий 2")
    montage.setOption("round_time", True)

    assert montage.config["layout_type"] == "Сценарий 2"
    assert montage.config["round_time"] is True
    assert "script2-container" in montage.html
    assert not bridge.project.canUndo
    saved = bridge._global_settings_service.load_settings()[
        "default_export_config"
    ]
    assert saved["layout_type"] == "Сценарий 2"
    assert saved["round_time"] is True


def test_qml_montage_accepts_table_column_width_changes():
    _app()
    bridge = AppBridge()

    bridge.montage.setOption("table_width_time", 6.5)
    bridge.montage.setOption("table_width_char", 12.0)
    bridge.montage.setOption("table_width_actor", 9.5)

    saved = bridge._global_settings_service.load_settings()[
        "default_export_config"
    ]
    assert saved["table_width_time"] == 6.5
    assert saved["table_width_char"] == 12.0
    assert saved["table_width_actor"] == 9.5


def test_qml_bridge_edits_montage_text_through_undo_stack(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.project_service.current_project_path = str(
        tmp_path / "editing.dub"
    )
    bridge._session.data["episodes"] = {"1": str(tmp_path / "one.ass")}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "modified_at": "before",
            "lines": [{
                "id": "stable-id",
                "start": 1.0,
                "end": 2.0,
                "character": "Hero",
                "text": "Original preview text",
            }],
        },
    }
    bridge.refresh()
    montage = bridge.montage
    montage.prepare("1")

    montage.updateText("0", "Edited in WebEngine")

    backups = list((tmp_path / ".backups").glob(
        "editing_editing_episode_1_*.dub_backup"
    ))
    assert len(backups) == 1
    line = bridge._session.data["episode_working_texts"]["1"]["lines"][0]
    assert line["text"] == "Edited in WebEngine"
    assert "Edited in WebEngine" in montage.html
    assert bridge.project.canUndo

    bridge.project.undo()
    assert line["text"] == "Original preview text"
    assert "Original preview text" in montage.html

    bridge.project.redo()
    assert line["text"] == "Edited in WebEngine"


def test_qml_bridge_keeps_actor_highlights_in_runtime_only():
    _app()
    bridge = AppBridge()
    bridge._session.data["actors"] = {
        "actor-1": {"name": "One", "color": "#112233"},
        "actor-2": {"name": "Two", "color": "#445566"},
    }
    bridge.refresh()

    montage = bridge.montage
    montage.setAllActorsHighlighted(False)
    montage.setActorHighlighted("actor-1", True)
    montage.setActorNegative("actor-1", True)
    rows = {row["actorId"]: row for row in montage.highlightModel.rows()}
    assert rows["actor-1"]["selected"] is True
    assert rows["actor-1"]["negative"] is True
    assert rows["actor-2"]["selected"] is False
    assert montage.highlightSummary == "1 из 2"

    montage.setActorHighlighted("actor-2", True)
    assert montage.config["highlight_ids_export"] is None
    assert montage.highlightSummary == "Все актёры"

    montage.setAllActorsHighlighted(False)
    assert montage.config["highlight_ids_export"] == []
    assert montage.highlightSummary == "Подсветка отключена"

    montage.setActorNegative("actor-2", True)
    assert montage.config["highlight_negative_ids_export"] == [
        "actor-1",
        "actor-2",
    ]
    assert not bridge.project.canUndo
    saved = bridge._global_settings_service.load_settings()[
        "default_export_config"
    ]
    assert "highlight_ids_export" not in saved
    assert "highlight_negative_ids_export" not in saved

    montage.reset()
    assert montage.config["highlight_ids_export"] is None
    assert montage.config["highlight_negative_ids_export"] == []


def test_qml_bridge_exports_montage_files_and_current_episode_batch(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["project_name"] = "Demo"
    bridge._session.data["episodes"] = {"1": str(tmp_path / "one.ass")}
    bridge._session.data["replica_merge_config"] = {"merge": False}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {
                    "id": "line-1",
                    "start": 1.0,
                    "end": 2.0,
                    "character": "Hero",
                    "text": "Export me",
                },
            ],
        },
    }
    bridge._session.data["export_config"]["open_auto"] = False
    bridge.refresh()
    montage = bridge.montage
    montage.prepare("1")

    html_path = tmp_path / "single"
    montage.exportFile("html", str(html_path))

    saved_html = html_path.with_suffix(".html")
    assert saved_html.exists()
    assert "Export me" in saved_html.read_text(encoding="utf-8")

    output_folder = tmp_path / "batch"
    montage.exportBatch(
        str(output_folder),
        True,
        False,
        True,
        False,
        False,
    )

    app = _app()
    for _ in range(100):
        app.processEvents()
        if not montage.batchBusy:
            break
    else:
        pytest.fail("Montage batch queue did not finish")

    assert (output_folder / "Demo - Ep1.html").exists()
    assert (output_folder / "Demo - Ep1.docx").exists()
    assert montage.batchCompleted == 2
    assert montage.batchResultModel.rowCount() == 2


def test_qml_montage_batch_reports_partial_failures_and_can_cancel(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["project_name"] = "Batch"
    bridge._session.data["episodes"] = {
        "1": str(tmp_path / "one.ass"),
        "2": str(tmp_path / "two.ass"),
    }
    bridge._session.data["episode_working_texts"] = {
        "1": {"lines": [{
            "id": "line-1", "start": 1.0, "end": 2.0,
            "character": "Hero", "text": "Available",
        }]},
    }
    bridge._session.data["export_config"]["open_auto"] = False
    bridge.refresh()
    montage = bridge.montage
    montage.prepare("1")

    output_folder = tmp_path / "partial"
    montage.exportBatch(
        str(output_folder), True, False, False, False, True
    )
    for _ in range(100):
        _app().processEvents()
        if not montage.batchBusy:
            break

    rows = montage.batchResultModel.rows()
    assert [row["statusKind"] for row in rows] == ["success", "error"]
    assert "ошибок: 1" in montage.batchSummary

    cancelled_folder = tmp_path / "cancelled"
    montage.exportBatch(
        str(cancelled_folder), True, False, True, False, True
    )
    montage.cancelBatch()
    for _ in range(100):
        _app().processEvents()
        if not montage.batchBusy:
            break
    assert montage.batchCompleted == 0
    assert not list(cancelled_folder.iterdir())


def test_qml_bridge_previews_and_exports_reaper_files(tmp_path):
    _app()
    bridge = AppBridge()
    video_path = tmp_path / "episode.mov"
    video_path.write_bytes(b"video")
    bridge._session.data["project_name"] = "Demo"
    bridge._session.data["episodes"] = {"1": str(tmp_path / "one.ass")}
    bridge._session.data["video_paths"] = {"1": str(video_path)}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"},
    }
    bridge._session.data["global_map"] = {"Hero": "actor-1"}
    bridge._session.data["replica_merge_config"] = {"merge": False}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "source": {"type": "srt"},
            "source_lines": [
                {
                    "id": "source-1",
                    "start": 1.0,
                    "end": 2.0,
                    "character": "Hero",
                    "text": "Source marker one",
                },
                {
                    "id": "source-2",
                    "start": 2.5,
                    "end": 3.5,
                    "character": "Hero",
                    "text": "Source marker two",
                },
            ],
            "lines": [
                {
                    "id": "line-1",
                    "start": 1.0,
                    "end": 3.5,
                    "character": "Hero",
                    "text": "Working replica",
                },
            ],
        },
    }
    bridge.refresh()

    reaper = bridge.reaper
    assert reaper.prepare() is True
    assert reaper.episode == "1"
    assert bridge.reaper.videoAvailable is True
    assert reaper.videoAvailable is True
    assert bridge.reaper.videoName == "episode.mov"
    assert bridge.reaper.sourceMarkersAvailable is True
    assert bridge.reaper.preview["regions"] == 1
    assert bridge.reaper.preview["tracks"] == 1
    assert bridge.reaper.preview["video"] is True

    source_preview = reaper.updatePreview(False, True, False, "source")
    assert source_preview["video"] is False
    assert source_preview["regions"] == 2
    assert bridge.reaper.preview["regions"] == 2
    assert "Source marker one" in bridge.reaper.preview["sample_regions"][0]

    rpp_path = tmp_path / "episode-project"
    assert reaper.export(
        "rpp", str(rpp_path), True, True, False, "source"
    ) is True
    saved_rpp = rpp_path.with_suffix(".rpp")
    assert saved_rpp.exists()
    assert "Source marker two" in saved_rpp.read_text(encoding="utf-8-sig")
    assert bridge.reaper.lastExportPath == str(saved_rpp)

    csv_path = tmp_path / "episode-markers"
    assert reaper.export(
        "csv", str(csv_path), False, True, False, "source"
    ) is True
    saved_csv = csv_path.with_suffix(".csv")
    assert saved_csv.exists()
    assert "Source marker one" in saved_csv.read_text(encoding="utf-8-sig")


def test_qml_bridge_reaper_export_requires_working_text(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": str(tmp_path / "one.ass")}
    bridge.refresh()

    assert bridge.reaper.prepare() is False
    assert bridge.reaper.preview == {}


def test_qml_bridge_exports_reaper_for_all_episodes(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["project_name"] = "Demo"
    bridge._session.data["episodes"] = {
        "1": str(tmp_path / "one.ass"),
        "2": str(tmp_path / "two.ass"),
    }
    bridge._session.data["video_paths"] = {}
    bridge._session.data["replica_merge_config"] = {"merge": False}
    bridge._session.data["episode_working_texts"] = {
        episode: {
            "source": {"type": "srt"},
            "source_lines": [{
                "id": f"source-{episode}",
                "start": 1.0,
                "end": 2.0,
                "character": "Hero",
                "text": f"Source episode {episode}",
            }],
            "lines": [{
                "id": f"line-{episode}",
                "start": 1.0,
                "end": 2.0,
                "character": "Hero",
                "text": f"Working episode {episode}",
            }],
        }
        for episode in ("1", "2")
    }
    bridge.refresh()

    reaper = bridge.reaper
    assert reaper.prepare()
    assert reaper.exportableEpisodeCount == 2
    assert reaper.allSourceMarkersAvailable
    assert not reaper.anyVideoAvailable

    rpp_folder = tmp_path / "rpp"
    assert reaper.exportAll(
        "rpp", str(rpp_folder), False, True, False, "source", "source_ass"
    )
    assert reaper.lastExportCount == 2
    assert reaper.lastExportPath == str(rpp_folder)
    assert "Source episode 1" in (
        rpp_folder / "Demo - Ep1.rpp"
    ).read_text(encoding="utf-8-sig")
    assert "Source episode 2" in (
        rpp_folder / "Demo - Ep2.rpp"
    ).read_text(encoding="utf-8-sig")

    csv_folder = tmp_path / "csv"
    assert reaper.exportAll(
        "csv", str(csv_folder), False, True, False, "merged", "project_episode"
    )
    assert (csv_folder / "Demo - 1.csv").exists()
    assert (csv_folder / "Demo - 2.csv").exists()


def test_qml_reaper_batch_rejects_partial_source_marker_mode(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {
        "1": str(tmp_path / "one.ass"),
        "2": str(tmp_path / "two.ass"),
    }
    bridge._session.data["video_paths"] = {}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "source": {"type": "srt"},
            "source_lines": [{
                "id": "source-1", "start": 1.0, "end": 2.0,
                "character": "Hero", "text": "Source",
            }],
            "lines": [{
                "id": "line-1", "start": 1.0, "end": 2.0,
                "character": "Hero", "text": "One",
            }],
        },
        "2": {"lines": [{
            "id": "line-2", "start": 2.0, "end": 3.0,
            "character": "Hero", "text": "Two",
        }]},
    }
    bridge.refresh()

    assert bridge.reaper.prepare()
    assert not bridge.reaper.allSourceMarkersAvailable
    assert not bridge.reaper.exportAll(
        "csv", str(tmp_path / "csv"), False, True, False, "source"
    )
    assert not (tmp_path / "csv").exists()


def test_qml_bridge_prepares_video_preview_and_filters_characters(tmp_path):
    _app()
    bridge = AppBridge()
    video_path = tmp_path / "episode.mp4"
    video_path.write_bytes(b"video")
    bridge._session.data["episodes"] = {"1": str(tmp_path / "one.ass")}
    bridge._session.data["video_paths"] = {"1": str(video_path)}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"},
    }
    bridge._session.data["global_map"] = {"Hero": "actor-1"}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {
                    "id": "line-1",
                    "start": 1.25,
                    "end": 2.5,
                    "character": "Hero",
                    "text": "Hero line",
                },
                {
                    "id": "line-2",
                    "start": 3.0,
                    "end": 4.0,
                    "character": "Villain",
                    "text": "Villain line",
                },
            ],
        },
    }
    bridge.refresh()

    video = bridge.video
    assert video.prepare("Hero") is True
    assert video.episode == "1"
    assert video.character == "Hero"
    assert video.hasVideo is True
    assert video.videoName == "episode.mp4"
    assert Path(video.videoUrl.toLocalFile()) == video_path
    assert video.count == 1
    assert [row["value"] for row in video.characterModel.rows()] == [
        "",
        "Hero",
        "Villain",
    ]
    preview_row = video.model.rows()[0]
    assert preview_row["startMs"] == 1250
    assert preview_row["endMs"] == 2500
    assert preview_row["actor"] == "Actor One"
    assert preview_row["actorColor"] == "#123456"

    video.setCharacter("")

    assert video.character == ""
    assert video.count == 2
    assert bridge.video.count == 2


def test_qml_bridge_video_preview_marks_multiple_assigned_actors(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": str(tmp_path / "one.ass")}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"},
        "actor-2": {"name": "Actor Two", "color": "#654321"},
    }
    bridge._session.data["global_map"] = {"Hero": ["actor-1", "actor-2"]}
    bridge._session.data["episode_working_texts"] = {
        "1": {"lines": [{
            "id": "line-1", "start": 1.25, "end": 2.5,
            "character": "Hero", "text": "Hero line",
        }]},
    }
    bridge.refresh()

    assert bridge.video.prepare("Hero") is True
    preview_row = bridge.video.model.rows()[0]

    assert preview_row["actor"] == "Несколько актёров"
    assert preview_row["actorColor"] == "transparent"


def test_qml_bridge_video_preview_keeps_replica_list_without_video(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": str(tmp_path / "one.ass")}
    bridge._session.data["video_paths"] = {"1": str(tmp_path / "missing.mp4")}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [{
                "id": "line-1",
                "start": 1.0,
                "end": 2.0,
                "character": "Hero",
                "text": "Replica only",
            }],
        },
    }
    bridge.refresh()

    video = bridge.video
    assert video.prepare("") is True
    assert video.hasVideo is False
    assert video.videoUrl.isEmpty()
    assert video.count == 1
    assert video.model.rows()[0]["text"] == "Replica only"


def test_qml_bridge_renames_episode_with_related_data(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": str(tmp_path / "episode.ass")}
    bridge._session.data["episode_actor_map"] = {"1": {"Hero": "actor-1"}}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {"id": "line-1", "start": 1.0, "end": 2.0, "character": "Hero", "text": "Alpha"},
            ]
        }
    }
    bridge.refresh()

    bridge.project.renameCurrentEpisode("Pilot")

    assert bridge.project.currentEpisode == "Pilot"
    assert "Pilot" in bridge._session.data["episodes"]
    assert "1" not in bridge._session.data["episodes"]
    assert bridge._session.data["episode_actor_map"]["Pilot"]["Hero"] == "actor-1"
    assert bridge._session.data["episode_working_texts"]["Pilot"]["lines"][0]["text"] == "Alpha"
    assert bridge.casting.charactersModel.rows()[0]["character"] == "Hero"

    bridge.project.undo()

    assert bridge.project.currentEpisode == "1"
    assert "1" in bridge._session.data["episodes"]
    assert "Pilot" not in bridge._session.data["episode_working_texts"]
    assert bridge._session.data["episode_working_texts"]["1"]["lines"][0]["text"] == "Alpha"


def test_qml_bridge_deletes_episode_with_undo(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": str(tmp_path / "episode.ass")}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {"id": "line-1", "start": 1.0, "end": 2.0, "character": "Hero", "text": "Alpha"},
            ]
        }
    }
    bridge.refresh()

    bridge.project.deleteCurrentEpisode()

    assert bridge.project.currentEpisode == ""
    assert bridge.project.episodesModel.rowCount() == 0
    assert "1" not in bridge._session.data["episodes"]
    assert "1" not in bridge._session.data["episode_working_texts"]

    bridge.project.undo()

    assert bridge.project.currentEpisode == "1"
    assert bridge.project.episodesModel.rowCount() == 1
    assert bridge.casting.charactersModel.rows()[0]["character"] == "Hero"


def test_qml_bridge_renames_character_with_undo(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": str(tmp_path / "episode.ass")}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"},
        "actor-2": {"name": "Actor Two", "color": "#654321"},
    }
    bridge._session.data["global_map"] = {"Hero": "actor-1"}
    bridge._session.data["episode_actor_map"] = {"1": {"Hero": "actor-2"}}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "characters": {"Hero": {"display_name": "Hero"}},
            "lines": [
                {
                    "id": "line-1",
                    "start": 1.0,
                    "end": 2.0,
                    "character": "Hero",
                    "display_character": "Hero",
                    "text": "Alpha words",
                },
            ],
        }
    }
    bridge.refresh()

    bridge.casting.renameCharacter("Hero", "Lead")

    payload = bridge._session.data["episode_working_texts"]["1"]
    assert bridge.casting.selectedCharacter == "Lead"
    assert bridge.casting.charactersModel.rows()[0]["character"] == "Lead"
    assert bridge._session.data["global_map"]["Lead"] == "actor-1"
    assert bridge._session.data["episode_actor_map"]["1"]["Lead"] == "actor-2"
    assert "Hero" not in bridge._session.data["global_map"]
    assert "Hero" not in bridge._session.data["episode_actor_map"]["1"]
    assert payload["lines"][0]["character"] == "Hero"
    assert payload["lines"][0]["display_character"] == "Lead"

    bridge.project.undo()

    payload = bridge._session.data["episode_working_texts"]["1"]
    assert bridge.casting.charactersModel.rows()[0]["character"] == "Hero"
    assert bridge._session.data["global_map"]["Hero"] == "actor-1"
    assert bridge._session.data["episode_actor_map"]["1"]["Hero"] == "actor-2"
    assert payload["lines"][0]["display_character"] == "Hero"


def test_qml_bridge_assign_actor_to_character_uses_undo_stack(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": str(tmp_path / "episode.ass")}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"},
    }
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {"id": "line-1", "start": 1.0, "end": 2.0, "character": "Hero", "text": "Alpha words"},
            ]
        }
    }
    bridge.refresh()

    bridge.casting.assignActor("Hero", "actor-1")

    assert bridge._session.data["global_map"]["Hero"] == "actor-1"
    assert bridge.casting.charactersModel.rows()[0]["actor"] == "Actor One"
    assert bridge.casting.charactersModel.rows()[0]["actorId"] == "actor-1"

    bridge.project.undo()
    assert "Hero" not in bridge._session.data["global_map"]
    assert bridge.casting.charactersModel.rows()[0]["actor"] == "-"

    bridge.project.redo()
    assert bridge._session.data["global_map"]["Hero"] == "actor-1"

    bridge.casting.assignActor("Hero", "")
    assert "Hero" not in bridge._session.data["global_map"]


def test_qml_bridge_assignment_scope_uses_episode_overrides(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {"1": str(tmp_path / "episode.ass")}
    bridge._session.data["actors"] = {
        "actor-1": {"name": "Actor One", "color": "#123456"},
        "actor-2": {"name": "Actor Two", "color": "#654321"},
    }
    bridge._session.data["global_map"] = {"Hero": "actor-1"}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {"id": "line-1", "start": 1.0, "end": 2.0, "character": "Hero", "text": "Alpha words"},
            ]
        }
    }
    bridge.refresh()

    assert bridge.casting.charactersModel.rows()[0]["scopeId"] == "global"
    assert bridge.casting.charactersModel.rows()[0]["actorId"] == "actor-1"

    bridge.casting.setAssignmentScope("Hero", "episode")

    assert bridge._session.data["episode_actor_map"]["1"]["Hero"] == "actor-1"
    assert bridge.casting.charactersModel.rows()[0]["scopeId"] == "episode"

    bridge.casting.assignActor("Hero", "actor-2")

    assert bridge._session.data["global_map"]["Hero"] == "actor-1"
    assert bridge._session.data["episode_actor_map"]["1"]["Hero"] == "actor-2"
    assert bridge.casting.charactersModel.rows()[0]["actorId"] == "actor-2"

    bridge.casting.setAssignmentScope("Hero", "global")

    assert "Hero" not in bridge._session.data["episode_actor_map"]["1"]
    assert bridge.casting.charactersModel.rows()[0]["scopeId"] == "global"
    assert bridge.casting.charactersModel.rows()[0]["actorId"] == "actor-1"


def test_qml_bridge_defers_destructive_action_until_user_decides():
    _app()
    bridge = AppBridge()
    project = bridge.project
    requested = []
    project.saveChangesRequested.connect(requested.append)
    project.name = "Unsaved work"

    project.create()

    assert project.name == "Unsaved work"
    assert requested == ["Сохранить изменения перед созданием нового проекта?"]

    project.resolvePendingChanges("cancel")
    assert project.name == "Unsaved work"

    project.create()
    project.resolvePendingChanges("discard")

    assert project.name == "Новый проект"
    assert not project.dirty


def test_qml_bridge_saves_unsaved_project_before_pending_action(tmp_path):
    _app()
    bridge = AppBridge()
    project = bridge.project
    save_path_requests = []
    project.savePathRequested.connect(lambda: save_path_requests.append(True))
    bridge.casting.addActor("Actor before reset")

    project.create()
    project.resolvePendingChanges("save")

    assert save_path_requests == [True]
    assert bridge.casting.actorsModel.rowCount() == 1

    saved_path = tmp_path / "before-reset.dub"
    project.saveAs(str(saved_path))

    assert saved_path.exists()
    assert project.name == "Новый проект"
    assert project.path == ""
    assert bridge.casting.actorsModel.rowCount() == 0


def test_qml_bridge_autosaves_dirty_saved_project(tmp_path):
    _app()
    bridge = AppBridge()
    project = bridge.project
    project_path = tmp_path / "autosave.dub"
    project.saveAs(str(project_path))
    project.name = "Changed"

    project.autoSave()

    backups = list((tmp_path / ".backups").glob("autosave_*.dub_backup"))
    assert len(backups) == 1
    assert project.dirty
    assert bridge.statusText == "Создана автокопия проекта"


def test_qml_ui_state_remembers_dialog_folders_and_window_values(tmp_path):
    _app()
    settings_path = tmp_path / "ui.ini"
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    source_file = source_dir / "episode.ass"
    source_file.write_text("", encoding="utf-8")

    settings = QSettings(
        str(settings_path),
        QSettings.Format.IniFormat,
    )
    state = UiStateBridge(settings)
    state.rememberFile("sourceFiles", str(source_file))
    state.setIntValue("main.width", 1440)
    state.setIntValue("teleprompterFloat.width", 428)
    state.setIntValue("teleprompterFloat.height", 516)
    state.setBoolValue("main.maximized", True)
    state.setBoolValue("actorColorCellFill", True)
    state.setStringValue("main.characterColumnsOrder", '["actor", "character"]')
    state.setIntValue("main.episodeTimelineColorMuteLevel", 1)
    state.setStringValue("main.episodeTimelineSortMode", "words")
    settings.sync()

    restored_settings = QSettings(
        str(settings_path),
        QSettings.Format.IniFormat,
    )
    restored = UiStateBridge(restored_settings)

    assert Path(restored.folderUrl("sourceFiles").toLocalFile()) == source_dir
    assert restored.intValue("main.width", 1000) == 1440
    assert restored.intValue("teleprompterFloat.width", 300) == 428
    assert restored.intValue("teleprompterFloat.height", 440) == 516
    assert restored.boolValue("main.maximized", False) is True
    assert restored.boolValue("actorColorCellFill", False) is True
    assert restored.stringValue("main.characterColumnsOrder", "") == '["actor", "character"]'
    assert restored.intValue("main.episodeTimelineColorMuteLevel", 2) == 1
    assert restored.stringValue("main.episodeTimelineSortMode", "") == "words"


def test_qml_bridge_restores_backup_after_unsaved_decision(tmp_path):
    _app()
    bridge = AppBridge()
    project = bridge.project
    project_path = tmp_path / "restorable.dub"
    requested = []
    project.saveChangesRequested.connect(requested.append)
    project.saveAs(str(project_path))
    project.name = "Версия из копии"
    project.autoSave()
    project.name = "Текущие несохранённые изменения"
    project.refreshBackups()

    rows = project.backupsModel.rows()
    assert len(rows) == 1
    assert rows[0]["modified"]
    assert rows[0]["size"].endswith("КБ")

    project.restoreBackup(rows[0]["path"])

    assert requested == [
        "Сохранить изменения перед восстановлением резервной копии?"
    ]
    assert project.name == "Текущие несохранённые изменения"

    project.resolvePendingChanges("discard")

    assert project.name == "Версия из копии"
    assert not project.dirty
    assert "Восстановлена резервная копия" in bridge.statusText
    assert list(
        (tmp_path / ".backups").glob(
            "restorable_before_restore_*.dub_backup"
        )
    )


def test_qml_bridge_builds_project_file_and_health_models(tmp_path):
    _app()
    bridge = AppBridge()
    missing_source = tmp_path / "missing.srt"
    bridge._session.data["episodes"] = {"1": str(missing_source)}
    bridge._session.data["episode_working_texts"] = {
        "1": {
            "lines": [
                {
                    "id": "1_0001",
                    "start": 1.0,
                    "end": 2.0,
                    "character": "Hero",
                    "text": "Hello",
                }
            ]
        }
    }

    bridge.refresh()

    project_files = bridge.projectFiles
    rows = project_files.filesModel.rows()
    assert [row["kind"] for row in rows] == ["source", "working", "video"]
    assert rows[0]["status"] == "Не найден"
    assert rows[1]["status"] == "Нет построчных реплик"
    assert project_files.currentEpisodeSourceMissing
    assert "предупреждения: 1" in project_files.healthSummary
    assert any(
        row["message"] == "Исходный файл серии не найден."
        for row in project_files.healthModel.rows()
    )


def test_qml_bridge_sets_folder_scans_links_and_undoes_atomically(tmp_path):
    _app()
    bridge = AppBridge()
    old_source = str(tmp_path / "elsewhere" / "Episode 1.srt")
    source = tmp_path / "Episode 1.srt"
    source.write_text(
        "1\n00:00:01,000 --> 00:00:02,000\nHero: Hello\n",
        encoding="utf-8",
    )
    bridge._session.data["episodes"] = {"1": old_source}
    bridge.refresh()

    project_files = bridge.projectFiles
    project_files.setFolder(str(tmp_path))

    assert project_files.folder == str(tmp_path)
    assert bridge._session.data["episodes"]["1"] == str(source)
    assert bridge.project.canUndo

    bridge.project.undo()

    assert project_files.folder == ""
    assert bridge._session.data["episodes"]["1"] == old_source


def test_qml_bridge_batch_imports_folder_sources_videos_and_texts(tmp_path):
    from docx import Document

    _app()
    bridge = AppBridge()
    srt_source = tmp_path / "Episode_01.srt"
    srt_source.write_text(
        "1\n00:00:01,000 --> 00:00:02,000\nHero: Subtitle\n",
        encoding="utf-8",
    )
    docx_source = tmp_path / "Episode_02.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Персонаж"
    table.cell(0, 1).text = "Текст"
    table.cell(1, 0).text = "Guest"
    table.cell(1, 1).text = "Document"
    document.save(docx_source)
    for name in ("Episode_01.mp4", "Episode_02.mov", "Episode_99.mp4"):
        (tmp_path / name).write_bytes(b"video")

    project_files = bridge.projectFiles
    project_files.setFolder(str(tmp_path))
    project_files.batchImportFolder()

    assert set(bridge._session.data["episodes"]) == {"1", "2"}
    assert set(bridge._session.data["video_paths"]) == {"1", "2"}
    assert set(bridge._session.data["episode_working_texts"]) == {"1", "2"}
    assert bridge._session.data["project_kind"] == "subtitle"

    bridge.project.undo()

    assert bridge._session.data["episodes"] == {}
    assert bridge._session.data["video_paths"] == {}
    assert bridge._session.data["episode_working_texts"] == {}
    assert project_files.folder == str(tmp_path)


def test_qml_bridge_removes_video_link_with_undo(tmp_path):
    _app()
    bridge = AppBridge()
    video = tmp_path / "Episode_01.mp4"
    video.write_bytes(b"video")
    bridge._session.data["episodes"] = {"1": str(tmp_path / "Episode_01.srt")}
    bridge._session.data["video_paths"] = {"1": str(video)}
    bridge.refresh()

    bridge.projectFiles.removeVideo("1")

    assert "1" not in bridge._session.data["video_paths"]
    assert bridge.project.canUndo

    bridge.project.undo()

    assert bridge._session.data["video_paths"]["1"] == str(video)


def test_qml_bridge_deletes_multiple_episodes_with_one_undo(tmp_path):
    _app()
    bridge = AppBridge()
    bridge._session.data["episodes"] = {
        "1": str(tmp_path / "Episode 1.srt"),
        "2": str(tmp_path / "Episode 2.srt"),
        "3": str(tmp_path / "Episode 3.srt"),
    }
    bridge._session.data["video_paths"] = {
        "1": str(tmp_path / "Episode 1.mp4"),
        "2": str(tmp_path / "Episode 2.mp4"),
    }
    bridge.project.selectEpisode("2")

    bridge.projectFiles.deleteEpisodes(["1", "2"])

    assert set(bridge._session.data["episodes"]) == {"3"}
    assert bridge._session.data["video_paths"] == {}
    assert bridge.project.currentEpisode == "3"

    bridge.project.undo()

    assert set(bridge._session.data["episodes"]) == {"1", "2", "3"}
    assert set(bridge._session.data["video_paths"]) == {"1", "2"}


def test_qml_bridge_relinks_source_and_undoes(tmp_path):
    _app()
    bridge = AppBridge()
    old_source = str(tmp_path / "missing.srt")
    new_source = tmp_path / "replacement.srt"
    new_source.write_text(
        "1\n00:00:01,000 --> 00:00:02,000\nHero: Hello\n",
        encoding="utf-8",
    )
    bridge._session.data["episodes"] = {"1": old_source}
    bridge.refresh()

    project_files = bridge.projectFiles
    project_files.relink("1", "source", str(new_source))

    assert bridge._session.data["episodes"]["1"] == str(new_source)
    assert not project_files.currentEpisodeSourceMissing

    bridge.project.undo()
    assert bridge._session.data["episodes"]["1"] == old_source


def test_qml_bridge_links_current_episode_video_and_undoes(tmp_path):
    _app()
    bridge = AppBridge()
    video = tmp_path / "Episode_01.mp4"
    video.write_bytes(b"video")
    bridge._session.data["episodes"] = {"1": str(tmp_path / "Episode_01.srt")}
    bridge.project.selectEpisode("1")

    bridge.projectFiles.relink("1", "video", str(video))

    assert bridge._session.data["video_paths"]["1"] == str(video)
    assert bridge.project.canUndo

    bridge.project.undo()
    assert "1" not in bridge._session.data["video_paths"]


def test_qml_bridge_regenerates_working_text_with_undo(tmp_path):
    _app()
    bridge = AppBridge()
    source = tmp_path / "Episode 1.srt"
    source.write_text(
        "1\n00:00:01,000 --> 00:00:02,000\nHero: Hello again\n",
        encoding="utf-8",
    )
    bridge._session.data["episodes"] = {"1": str(source)}
    bridge.refresh()

    bridge.projectFiles.regenerateWorkingText("1")

    assert bridge._session.data["episode_working_texts"]["1"]["lines"][0]["text"] == "Hello again"
    assert bridge.casting.charactersModel.rows()[0]["character"] == "Hero"

    bridge.project.undo()

    assert "1" not in bridge._session.data["episode_working_texts"]
    assert bridge.casting.charactersModel.rowCount() == 0


def test_qml_bridge_regenerates_docx_using_saved_mapping_without_merging(tmp_path):
    from docx import Document

    _app()
    bridge = AppBridge()
    source = tmp_path / "Episode 2.docx"
    document = Document()
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Реплика"
    table.cell(0, 1).text = "Кто"
    table.cell(1, 0).text = "First"
    table.cell(1, 1).text = "Hero"
    table.cell(2, 0).text = "Second"
    table.cell(2, 1).text = "Hero"
    document.save(source)
    bridge._session.data["episodes"] = {"2": str(source)}
    bridge._global_settings_service.update_docx_import_config({
        "header_mode": "first",
        "minimum_header_matches": 1,
        "mapping": {
            "character": 1,
            "time_start": None,
            "time_end": None,
            "time_split": None,
            "text": 0,
        },
    })
    bridge.refresh()

    bridge.projectFiles.regenerateWorkingText("2")

    payload = bridge._session.data["episode_working_texts"]["2"]
    assert [line["text"] for line in payload["lines"]] == ["First", "Second"]
    assert [line["character"] for line in payload["lines"]] == ["Hero", "Hero"]
    assert any(
        row["kind"] == "working" and row["canRegenerate"]
        for row in bridge.projectFiles.filesModel.rows()
    )

    bridge.project.undo()

    assert "2" not in bridge._session.data["episode_working_texts"]


def test_qml_bridge_creates_missing_srt_and_docx_texts_with_one_undo(tmp_path):
    from docx import Document

    _app()
    bridge = AppBridge()
    srt_source = tmp_path / "Episode 1.srt"
    srt_source.write_text(
        "1\n00:00:01,000 --> 00:00:02,000\nHero: Subtitle\n",
        encoding="utf-8",
    )
    docx_source = tmp_path / "Episode 2.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Персонаж"
    table.cell(0, 1).text = "Текст"
    table.cell(1, 0).text = "Guest"
    table.cell(1, 1).text = "Document line"
    document.save(docx_source)
    bridge._session.data["episodes"] = {
        "1": str(srt_source),
        "2": str(docx_source),
    }
    bridge._global_settings_service.update_docx_import_config({"header_mode": "first"})
    bridge.refresh()

    bridge.projectFiles.createMissingWorkingTexts()

    assert set(bridge._session.data["episode_working_texts"]) == {"1", "2"}
    assert bridge._session.data["episode_working_texts"]["2"]["lines"][0][
        "text"
    ] == "Document line"
    assert bridge.project.canUndo

    bridge.project.undo()

    assert bridge._session.data["episode_working_texts"] == {}
    assert not bridge.project.canUndo
