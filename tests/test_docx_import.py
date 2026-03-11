"""Тесты для импорта DOCX файлов"""

import pytest
import os
import tempfile
from services.docx_import_service import DocxImportService, DEFAULT_COLUMN_MAPPING, DEFAULT_TIME_SEPARATORS

try:
    from docx import Document
    from docx.table import Table
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
class TestDocxImportService:
    """Тесты для сервиса импорта DOCX"""

    def setup_method(self):
        """Настройка перед каждым тестом"""
        self.service = DocxImportService()

    def _create_test_docx(self, rows):
        """Создание тестового DOCX файла"""
        temp_file = tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False, mode='wb'
        )
        temp_path = temp_file.name
        temp_file.close()

        doc = Document()
        table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows else 0)

        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                table.cell(i, j).text = cell_text

        doc.save(temp_path)
        return temp_path

    def test_extract_tables_from_docx(self):
        """Тест извлечения таблиц из DOCX"""
        rows = [
            ["Персонаж", "Начало", "Конец", "Текст"],
            ["Иван", "00:00:01,000", "00:00:03,000", "Привет!"],
            ["Мария", "00:00:04,000", "00:00:06,000", "Как дела?"]
        ]

        temp_path = self._create_test_docx(rows)

        try:
            all_tables = self.service.extract_tables_from_docx(temp_path)
            assert len(all_tables) == 1  # Одна таблица
            extracted = all_tables[0]
            assert len(extracted) == 3
            assert extracted[0] == ["Персонаж", "Начало", "Конец", "Текст"]
        finally:
            os.unlink(temp_path)

    def test_extract_multiple_tables(self):
        """Тест извлечения нескольких таблиц"""
        # Создаём документ с двумя таблицами
        temp_file = tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False, mode='wb'
        )
        temp_path = temp_file.name
        temp_file.close()

        doc = Document()
        
        # Первая таблица
        table1 = doc.add_table(rows=2, cols=3)
        table1.cell(0, 0).text = "Персонаж"
        table1.cell(0, 1).text = "Тайминг"
        table1.cell(0, 2).text = "Текст"
        table1.cell(1, 0).text = "Иван"
        table1.cell(1, 1).text = "00:00:01,000 - 00:00:03,000"
        table1.cell(1, 2).text = "Привет!"
        
        # Добавляем параграф между таблицами
        doc.add_paragraph()
        
        # Вторая таблица
        table2 = doc.add_table(rows=2, cols=3)
        table2.cell(0, 0).text = "Персонаж"
        table2.cell(0, 1).text = "Тайминг"
        table2.cell(0, 2).text = "Текст"
        table2.cell(1, 0).text = "Мария"
        table2.cell(1, 1).text = "00:00:04,000 - 00:00:06,000"
        table2.cell(1, 2).text = "Как дела?"

        doc.save(temp_path)

        try:
            all_tables = self.service.extract_tables_from_docx(temp_path)
            assert len(all_tables) == 2  # Две таблицы
            assert len(all_tables[0]) == 2
            assert len(all_tables[1]) == 2
            assert all_tables[0][1][0] == "Иван"
            assert all_tables[1][1][0] == "Мария"
        finally:
            os.unlink(temp_path)

    def test_detect_columns_auto(self):
        """Тест автоопределения колонок"""
        rows = [
            ["Персонаж", "Тайминг", "Текст"],
            ["Иван", "00:00:01,000 - 00:00:03,000", "Привет!"]
        ]

        mapping = self.service.detect_columns(rows)
        assert mapping['character'] == 0
        assert mapping['time_split'] == 1
        assert mapping['text'] == 2
        assert self.service._has_header == True

    def test_detect_columns_split_timing(self):
        """Тест автоопределения для тайминга в одной колонке"""
        rows = [
            ["Персонаж", "Тайминг", "Текст"],
            ["Иван", "00:00:01,000 - 00:00:03,000", "Привет!"]
        ]

        mapping = self.service.detect_columns(rows)
        assert mapping['character'] == 0
        assert mapping['time_split'] == 1
        assert mapping['text'] == 2
        assert self.service._has_header == True

    def test_detect_no_header(self):
        """Тест определения без заголовка"""
        rows = [
            ["Иван", "00:00:01,000 - 00:00:03,000", "Привет!"],
            ["Мария", "00:00:04,000 - 00:00:06,000", "Как дела?"]
        ]

        mapping = self.service.detect_columns(rows)
        # Заголовок не найден, используем значения по умолчанию
        assert mapping['character'] == 0
        assert mapping['time_split'] == 1  # По умолчанию time_split в колонке 1
        assert mapping['text'] == 2
        assert self.service._has_header == False

    def test_parse_with_mapping(self):
        """Тест парсинга с заданным маппингом"""
        rows = [
            ["Персонаж", "Тайминг", "Текст"],
            ["Иван", "00:00:01,000 - 00:00:03,000", "Привет!"],
            ["Мария", "00:00:04,000 - 00:00:06,000", "Как дела?"]
        ]

        mapping = {
            'character': 0,
            'time_start': None,
            'time_end': None,
            'time_split': 1,
            'text': 2
        }

        self.service._has_header = True
        stats, lines = self.service.parse_with_mapping(rows, mapping)

        assert len(lines) == 2
        assert lines[0]['char'] == "Иван"
        assert lines[0]['text'] == "Привет!"
        assert abs(lines[0]['s'] - 1.0) < 0.01
        assert abs(lines[0]['e'] - 3.0) < 0.01

        assert len(stats) == 2
        assert stats[0]['name'] == "Иван"
        assert stats[0]['lines'] == 1
        assert stats[0]['rings'] == 1  # Для DOCX rings = lines

    def test_parse_split_time(self):
        """Тест парсинга тайминга в одной колонке"""
        # Тест с разделителем по умолчанию '-'
        test_cases = [
            ("00:00:01,000 - 00:00:03,000", 1.0, 3.0),
            ("00:00:01,000-00:00:03,000", 1.0, 3.0),  # без пробелов
        ]

        for time_str, expected_start, expected_end in test_cases:
            start, end = self.service._parse_split_time(time_str)
            assert abs(start - expected_start) < 0.01, f"Failed start for {time_str}"
            assert abs(end - expected_end) < 0.01, f"Failed end for {time_str}"

    def test_parse_split_time_multiple_separators(self):
        """Тест парсинга с несколькими разделителями"""
        service = DocxImportService(time_separators=['-', '–', '—', '|', '/'])
        
        test_cases = [
            ("00:00:01,000 - 00:00:03,000", 1.0, 3.0),
            ("00:00:01,000–00:00:03,000", 1.0, 3.0),  # en-dash
            ("00:00:01,000—00:00:03,000", 1.0, 3.0),  # em-dash
            ("00:00:01,000 | 00:00:03,000", 1.0, 3.0),
            ("00:00:01,000 / 00:00:03,000", 1.0, 3.0),
        ]

        for time_str, expected_start, expected_end in test_cases:
            start, end = service._parse_split_time(time_str)
            assert abs(start - expected_start) < 0.01, f"Failed start for {time_str}"
            assert abs(end - expected_end) < 0.01, f"Failed end for {time_str}"

    def test_parse_split_time_with_custom_separators(self):
        """Тест парсинга с кастомными разделителями"""
        service = DocxImportService(time_separators=['::', '->'])
        
        start, end = service._parse_split_time("00:00:01,000 :: 00:00:03,000")
        assert abs(start - 1.0) < 0.01
        assert abs(end - 3.0) < 0.01
        
        start, end = service._parse_split_time("00:00:01,000 -> 00:00:03,000")
        assert abs(start - 1.0) < 0.01
        assert abs(end - 3.0) < 0.01

    def test_parse_with_split_timing(self):
        """Тест парсинга с таймингом в одной колонке"""
        rows = [
            ["Персонаж", "Тайминг", "Текст"],
            ["Иван", "00:00:01,000 - 00:00:03,000", "Привет!"],
            ["Мария", "00:00:04,000 - 00:00:06,000", "Как дела?"]
        ]

        mapping = {
            'character': 0,
            'time_start': None,
            'time_end': None,
            'time_split': 1,
            'text': 2
        }

        self.service._has_header = True
        stats, lines = self.service.parse_with_mapping(rows, mapping)

        assert len(lines) == 2
        assert lines[0]['char'] == "Иван"
        assert lines[0]['text'] == "Привет!"
        assert abs(lines[0]['s'] - 1.0) < 0.01
        assert abs(lines[0]['e'] - 3.0) < 0.01

    def test_parse_time_formats(self):
        """Тест парсинга различных форматов времени"""
        test_cases = [
            ("00:00:01,000", 1.0),
            ("00:00:01.500", 1.5),
            ("00:01:30,000", 90.0),
            ("01:00:00,000", 3600.0),
            ("00:30", 30.0),
            ("1:30", 90.0),
        ]

        for time_str, expected in test_cases:
            result = self.service._parse_time(time_str)
            assert abs(result - expected) < 0.01, f"Failed for {time_str}"

    def test_get_preview_data(self):
        """Тест получения данных для предпросмотра"""
        rows = [
            ["Персонаж", "Тайминг", "Текст"],
            ["Иван", "00:00:01,000 - 00:00:03,000", "Привет!"],
            ["Мария", "00:00:04,000 - 00:00:06,000", "Как дела?"]
        ]

        mapping = {
            'character': 0,
            'time_start': None,
            'time_end': None,
            'time_split': 1,
            'text': 2
        }

        self.service._has_header = True
        preview = self.service.get_preview_data(rows, mapping, limit=2)
        assert len(preview) == 2
        assert preview[0]['mapped']['character'] == "Иван"
        assert preview[0]['mapped']['text'] == "Привет!"
        assert preview[0]['time_split_start_parsed'] == 1.0
        assert preview[0]['time_split_end_parsed'] == 3.0

    def test_get_preview_data_split_time(self):
        """Тест получения данных для предпросмотра с split таймингом"""
        rows = [
            ["Персонаж", "Тайминг", "Текст"],
            ["Иван", "00:00:01,000 - 00:00:03,000", "Привет!"],
        ]

        mapping = {
            'character': 0,
            'time_start': None,
            'time_end': None,
            'time_split': 1,
            'text': 2
        }

        self.service._has_header = True
        preview = self.service.get_preview_data(rows, mapping, limit=1)
        assert len(preview) == 1
        assert preview[0]['time_split_start_parsed'] == 1.0
        assert preview[0]['time_split_end_parsed'] == 3.0

    def test_empty_file(self):
        """Тест обработки пустого файла"""
        rows = []
        mapping = DEFAULT_COLUMN_MAPPING.copy()

        stats, lines = self.service.parse_with_mapping(rows, mapping)
        assert len(lines) == 0
        assert len(stats) == 0

    def test_missing_columns(self):
        """Тест обработки файла с отсутствующими колонками"""
        rows = [
            ["Данные"],
            ["Привет!"],
            ["Как дела?"]
        ]

        mapping = {
            'character': None,
            'time_start': None,
            'time_end': None,
            'time_split': None,
            'text': 0
        }

        self.service._has_header = False
        stats, lines = self.service.parse_with_mapping(rows, mapping)
        assert len(lines) == 3  # Все строки включая первую
        assert lines[0]['char'] == ""
        assert lines[0]['text'] == "Данные"  # Первая строка тоже парсится

    def test_parse_without_header(self):
        """Тест парсинга без заголовка (первая строка не пропускается)"""
        rows = [
            ["Иван", "00:00:01,000 - 00:00:03,000", "Привет!"],
            ["Мария", "00:00:04,000 - 00:00:06,000", "Как дела?"]
        ]

        mapping = {
            'character': 0,
            'time_start': None,
            'time_end': None,
            'time_split': 1,
            'text': 2
        }

        self.service._has_header = False
        stats, lines = self.service.parse_with_mapping(rows, mapping)

        assert len(lines) == 2  # Обе строки включая первую
        assert lines[0]['char'] == "Иван"
        assert lines[0]['text'] == "Привет!"
