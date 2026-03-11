"""Тесты для docx_import_service.py"""

import pytest
import tempfile
import os
from unittest.mock import patch, MagicMock

from services.docx_import_service import (
    DocxImportService,
    DOCX_AVAILABLE,
    DEFAULT_COLUMN_MAPPING,
    DEFAULT_TIME_SEPARATORS,
)


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
class TestDocxImportService:
    """Тесты для DocxImportService"""

    @pytest.fixture
    def service(self):
        """Сервис для тестов"""
        return DocxImportService()

    @pytest.fixture
    def sample_docx_path(self, tmp_path):
        """Путь к временному DOCX файлу"""
        from docx import Document
        
        doc_path = tmp_path / "test.docx"
        doc = Document()
        
        # Добавляем таблицу
        table = doc.add_table(rows=4, cols=3)
        table.cell(0, 0).text = "Персонаж"
        table.cell(0, 1).text = "Тайминг"
        table.cell(0, 2).text = "Текст"
        
        table.cell(1, 0).text = "Иван"
        table.cell(1, 1).text = "00:00:01,000 - 00:00:03,000"
        table.cell(1, 2).text = "Привет!"
        
        table.cell(2, 0).text = "Мария"
        table.cell(2, 1).text = "00:00:04,000 - 00:00:06,000"
        table.cell(2, 2).text = "Как дела?"
        
        table.cell(3, 0).text = "Иван"
        table.cell(3, 1).text = "00:00:07,000 - 00:00:09,000"
        table.cell(3, 2).text = "Нормально."
        
        doc.save(str(doc_path))
        return str(doc_path)

    def test_init_default_values(self, service):
        """Тест значений по умолчанию"""
        assert service.merge_gap == 5
        assert service.fps == 25.0
        assert service.time_separators == DEFAULT_TIME_SEPARATORS

    def test_set_merge_gap(self, service):
        """Тест установки зазора"""
        service.set_merge_gap(10)
        assert service.merge_gap == 10

    def test_set_fps(self, service):
        """Тест установки FPS"""
        service.set_fps(30.0)
        assert service.fps == 30.0

    def test_set_time_separators(self, service):
        """Тест установки разделителей"""
        service.set_time_separators(['|', '/'])
        assert service.time_separators == ['|', '/']

    def test_extract_tables_from_docx(self, sample_docx_path, service):
        """Тест извлечения таблиц"""
        tables = service.extract_tables_from_docx(sample_docx_path)
        
        assert len(tables) >= 1
        assert len(tables[0]) == 4  # 4 строки
        assert len(tables[0][0]) == 3  # 3 колонки

    def test_extract_first_table(self, sample_docx_path, service):
        """Тест извлечения первой таблицы"""
        table = service.extract_first_table(sample_docx_path)
        
        assert len(table) == 4
        assert table[0][0] == "Персонаж"

    def test_extract_tables_no_tables(self, tmp_path, service):
        """Тест когда нет таблиц"""
        from docx import Document
        
        doc_path = tmp_path / "empty.docx"
        doc = Document()
        doc.add_paragraph("Just text")
        doc.save(str(doc_path))
        
        tables = service.extract_tables_from_docx(str(doc_path))
        
        # Должен вернуть текст как таблицу
        assert len(tables) >= 1

    def test_detect_columns(self, sample_docx_path, service):
        """Тест определения колонок"""
        table = service.extract_first_table(sample_docx_path)
        mapping = service.detect_columns(table)
        
        assert 'character' in mapping
        assert 'text' in mapping
        assert 'time_split' in mapping

    def test_detect_columns_no_header(self, service):
        """Тест определения без заголовка"""
        rows = [
            ["Иван", "00:00:01,000", "Привет"],
            ["Мария", "00:00:02,000", "Пока"],
        ]
        
        mapping = service.detect_columns(rows)
        
        assert mapping is not None

    def test_detect_columns_empty(self, service):
        """Тест определения для пустых данных"""
        mapping = service.detect_columns([])
        
        assert mapping == DEFAULT_COLUMN_MAPPING.copy()

    def test_get_available_columns(self, service):
        """Тест получения доступных колонок"""
        rows = [
            ["a", "b", "c"],
            ["1", "2", "3"],
        ]
        
        columns = service.get_available_columns(rows)
        
        assert columns == [0, 1, 2]

    def test_get_available_columns_empty(self, service):
        """Тест для пустых данных"""
        columns = service.get_available_columns([])
        
        assert columns == []

    def test_parse_with_mapping(self, sample_docx_path, service):
        """Тест парсинга с маппингом"""
        table = service.extract_first_table(sample_docx_path)
        mapping = service.detect_columns(table)
        
        stats, lines = service.parse_with_mapping(table, mapping)
        
        assert len(lines) == 3  # 3 реплики (без заголовка)
        assert len(stats) == 2  # 2 персонажа
        
        # Проверяем статистику
        ivan_stats = next((s for s in stats if s['name'] == 'Иван'), None)
        assert ivan_stats is not None
        assert ivan_stats['lines'] == 2

    def test_parse_with_mapping_empty_rows(self, service):
        """Тест парсинга пустых строк"""
        rows = [
            ["", "", ""],
            ["", "", ""],
        ]
        mapping = DEFAULT_COLUMN_MAPPING.copy()
        
        stats, lines = service.parse_with_mapping(rows, mapping)
        
        assert len(lines) == 0

    def test_parse_with_mapping_missing_columns(self, service):
        """Тест парсинга с отсутствующими колонками"""
        rows = [
            ["Иван", "Привет"],  # Нет тайминга
        ]
        mapping = {
            'character': 0,
            'time_start': None,
            'time_end': None,
            'time_split': None,
            'text': 1,
        }
        
        stats, lines = service.parse_with_mapping(rows, mapping)
        
        assert len(lines) == 1
        assert lines[0]['s'] == 0.0  # Значение по умолчанию

    def test_get_cell_value(self, service):
        """Тест получения значения ячейки"""
        row = ["a", "b", "c"]
        
        assert service._get_cell_value(row, 0) == "a"
        assert service._get_cell_value(row, 1) == "b"
        assert service._get_cell_value(row, 5) is None  # За пределами
        assert service._get_cell_value(row, None) is None

    def test_parse_time_hh_mm_ss_ms(self, service):
        """Тест парсинга HH:MM:SS,mmm"""
        result = service._parse_time("00:01:02,500")
        assert result == 62.5

    def test_parse_time_mm_ss_ms(self, service):
        """Тест парсинга MM:SS,mmm"""
        result = service._parse_time("01:30,500")
        assert result == 90.5

    def test_parse_time_hh_mm_ss(self, service):
        """Тест парсинга HH:MM:SS"""
        result = service._parse_time("00:01:30")
        assert result == 90.0

    def test_parse_time_mm_ss(self, service):
        """Тест парсинга MM:SS"""
        result = service._parse_time("02:30")
        assert result == 150.0

    def test_parse_time_none(self, service):
        """Тест парсинга None"""
        assert service._parse_time(None) is None
        assert service._parse_time("") is None

    def test_parse_time_invalid(self, service):
        """Тест парсинга невалидного"""
        result = service._parse_time("invalid")
        # Функция возвращает 0.0 при ошибке парсинга
        assert isinstance(result, float)

    def test_parse_split_time(self, service):
        """Тест парсинга разделённого тайминга"""
        start, end = service._parse_split_time("00:00:01,000 - 00:00:03,000")
        
        assert start == 1.0
        assert end == 3.0

    def test_parse_split_time_pipe_separator(self, service):
        """Тест парсинга с разделителем |"""
        service.set_time_separators(['|'])
        start, end = service._parse_split_time("00:00:01,000|00:00:03,000")
        
        assert start == 1.0
        assert end == 3.0

    def test_parse_split_time_none(self, service):
        """Тест парсинга None"""
        start, end = service._parse_split_time(None)
        
        assert start is None
        assert end is None

    def test_parse_split_time_no_separator(self, service):
        """Тест без разделителя"""
        start, end = service._parse_split_time("00:00:01,000")
        
        assert start is None
        assert end is None

    def test_get_preview_data(self, sample_docx_path, service):
        """Тест получения данных для предпросмотра"""
        table = service.extract_first_table(sample_docx_path)
        mapping = service.detect_columns(table)
        
        preview = service.get_preview_data(table, mapping, limit=2)
        
        assert len(preview) == 2
        assert 'raw' in preview[0]
        assert 'mapped' in preview[0]

    def test_get_preview_data_with_header(self, sample_docx_path, service):
        """Тест предпросмотра с заголовком"""
        table = service.extract_first_table(sample_docx_path)
        mapping = service.detect_columns(table)
        
        preview = service.get_preview_data(table, mapping, limit=10)
        
        # Первая строка - заголовок, должна быть пропущена
        assert len(preview) == 3  # 3 строки данных

    def test_docx_not_available(self, tmp_path):
        """Тест когда docx не доступен"""
        with patch('services.docx_import_service.DOCX_AVAILABLE', False):
            service = DocxImportService()
            tables = service.extract_tables_from_docx(str(tmp_path / "test.docx"))
            
            assert tables == []

    def test_extract_tables_error(self, service):
        """Тест ошибки извлечения"""
        with patch('services.docx_import_service.Document', side_effect=Exception("Error")):
            tables = service.extract_tables_from_docx("/nonexistent.docx")
            
            assert tables == []

    def test_parse_with_mapping_error(self, service):
        """Тест ошибки парсинга"""
        rows = [["test"]]
        mapping = DEFAULT_COLUMN_MAPPING.copy()
        
        # Не должно вызывать исключение
        stats, lines = service.parse_with_mapping(rows, mapping)


class TestDocxImportServiceNoDocx:
    """Тесты когда python-docx не установлен"""

    def test_docx_not_available(self):
        """Тест флага доступности"""
        # Если docx не доступен, сервис должен работать с ограничениями
        if not DOCX_AVAILABLE:
            service = DocxImportService()
            assert service is not None
